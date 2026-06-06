"""文件解析服务 —— 支持 DOCX、PDF、TXT 格式的病历文件文本提取"""

import io


class FileParseResult:
    def __init__(self, text: str, filename: str, file_type: str, page_count: int = 1,
                 parse_time_ms: float = 0.0):
        self.text = text
        self.filename = filename
        self.file_type = file_type
        self.page_count = page_count
        self.parse_time_ms = parse_time_ms


async def parse_file(content: bytes, filename: str) -> FileParseResult:
    """根据文件扩展名自动选择解析器"""
    import time
    t0 = time.time()

    lower = filename.lower()
    if lower.endswith(".txt"):
        text = content.decode("utf-8", errors="replace")
        file_type = "txt"
        pages = 1
    elif lower.endswith(".docx"):
        text, pages = _parse_docx(content)
        file_type = "docx"
    elif lower.endswith(".pdf"):
        text, pages = _parse_pdf(content)
        file_type = "pdf"
    else:
        raise ValueError(f"不支持的文件格式: {filename}")

    elapsed = (time.time() - t0) * 1000
    return FileParseResult(
        text=text, filename=filename, file_type=file_type,
        page_count=pages, parse_time_ms=elapsed,
    )


def _parse_docx(content: bytes) -> tuple[str, int]:
    """解析DOCX文件，返回(文本内容, 段落数)"""
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs), len(paragraphs)


def _parse_pdf(content: bytes) -> tuple[str, int]:
    """解析PDF文件，返回(文本内容, 页数)"""
    from pypdf import PdfReader  # pypdf is the actively-maintained successor to PyPDF2 (drop-in compatible API)

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            pages.append(text.strip())

    return "\n\n".join(pages), len(reader.pages)
