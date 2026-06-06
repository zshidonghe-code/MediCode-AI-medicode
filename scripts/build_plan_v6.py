# coding: utf-8
"""Build merged three-week plan v6 .docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

sty = doc.styles['Normal']
sty.font.name = '微软雅黑'; sty.font.size = Pt(10.5)
sty.paragraph_format.space_after = Pt(4); sty.paragraph_format.line_spacing = 1.2
sty.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for lv in range(1,4):
    h = doc.styles['Heading %d'%lv]
    h.font.name = '微软雅黑'; h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(0x0E,0xA5,0xE9); h.font.bold = True

BLUE = RGBColor(0x0E,0xA5,0xE9); DARK = RGBColor(0x1E,0x29,0x3B)
GRAY = RGBColor(0x94,0xA3,0xB8); RED = RGBColor(0xEF,0x44,0x44); GREEN = RGBColor(0x10,0xB9,0x81)

def P(text, bold=False, size=None, color=None, align=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.left_indent = Cm(indent)
    return p

def H1(text): return doc.add_heading(text, level=1)
def H2(text): return doc.add_heading(text, level=2)
def B(text, indent=0.5): P('• ' + text, size=10.5, color=DARK, indent=indent)

def TABLE(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(hd); r.bold = True; r.font.size = Pt(9)
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val)); r.font.size = Pt(9)
            r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

def DIVIDER():
    P('━' * 60, size=6, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# === COVER ===
P('三周综合计划', bold=True, size=26, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
P('自学 · 乒乓球力量训练 · 项目工作 · 转专业冲刺', size=13, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
P('完美主义的解药不是更努力，是允许自己做得不够好。', size=11, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
P('郑诗东和 | 上海对外经贸大学 | 2026/5/29 - 6/18 | v6 合并版', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
P('', size=6)

# === 一 ===
H1('一、核心原则')
P('以下原则高于一切具体安排。当计划细节与原则冲突时，以原则为准。', bold=True)
B('产出重于时间：每天完成两张可验证的任务卡就算赢。不计算学了多少小时。')
B('粗糙执行 > 精确放弃：一份执行了60%的粗计划，比执行了0%的完美计划好一万倍。')
B('降载不是备用计划：状态差的日子，启动降载。只做一道微积分题 + 15分钟任何运动 = 赢了。')
B('微积分是底线：唯一与转专业硬门槛(80分+)直接挂钩的科目。崩盘时其他全可放，微积分不能放。')
P('', size=4)

# === 二 ===
H1('二、v6 改进说明（v3 + v5 合并优化）')
B('统一每日节奏：仅周二Python课、周四体育课必须到场，其余全部自学。统一9:00起床。')
B('码医投入翻倍：从仅周六 -> 周二/四晚各1.5h + 周六3h = 每周6h。码医是面试核心素材。')
B('思政优先级上调：从P3 -> P2。考前突击收益极高，背了就有分。英语从P2下调至P3。')
B('手机物理隔离：22:30手机移出卧室，用实体闹钟起床。不是意志力问题，是手机在床边的问题。')
B('小艺+闹钟双保险：AI助理小艺Claw负责每日推送追踪，手机闹钟作为离线备份。')
P('', size=4)

# === 三 ===
H1('三、每日节奏（统一版）')
P('周二 Python(10:00-11:40) 和周四体育(14:45-16:20)为仅有的两门必须到课，其余全部自学。', size=10, color=GRAY)
P('', size=2)
TABLE(['时间', '内容', '类型'], [
    ['9:00-9:30', '起床、洗漱、早餐。看一眼今天两张任务卡。', '启动'],
    ['9:30-10:00', '通勤/预热：回顾昨日薄弱点，心里过一遍今天的任务卡。', '缓冲'],
    ['10:00-12:00', '任务卡#1：数学（微积分B 或 概率论B）。大脑最清醒，攻克最难内容。', '核心'],
    ['12:00-13:00', '午饭+休息。不碰学习，让大脑真正停下来。', '休息'],
    ['13:00-15:00', '任务卡#2：编程/微观经济学/英语/码医。下午灵活安排。', '核心'],
    ['15:00-16:30', '训练日：乒乓球或力量训练。（周四体育课已覆盖）', '身体'],
    ['16:30-18:00', '弹性：推进未完成任务卡/复习薄弱点/码医/股票复盘。补漏优先。', '弹性'],
    ['18:00-19:00', '晚饭+休息。', '休息'],
    ['19:00-21:00', '晚间自习：巩固今日+预习明天。重点收尾补漏。周二/四晚间优先排码医1.5h。', '核心'],
    ['21:00-21:30', '向小艺汇报：今天任务卡完成了吗？一句话记进度。', '收尾'],
    ['21:30-22:30', '不可支配自由时间。社交、刷手机、看书、发呆。是权利，不是奖励。', '自由'],
    ['22:30', '手机移出卧室。设好实体闹钟。', '硬规则'],
    ['23:00', '熄灯。睡眠是第一生产力。', '硬规则'],
])
P('', size=4)

# === 四 ===
H1('四、每周任务卡矩阵')
P('每天两张卡，完成即赢。周二Python(10:00到课)、周四体育(14:45到课)、周二/四/六码医。', size=10, color=GRAY)
P('', size=2)
TABLE(['', '任务卡#1（数学）', '任务卡#2（编程/英语/码医）', '训练', '备注'], [
    ['周一', '微积分B 课后习题+整理错题', '英语视听说 听力1篇+背单词', '力量A 腿+核心', ''],
    ['周二', '概率论B 当前章节课后习题', 'Python到课(10:00)+作业\n晚间：码医1.5h', '乒乓球90min', '码医优先'],
    ['周三', '微积分B 课后习题+整理错题', '微观经济学 读教材+画供需曲线', '力量B 上肢+旋转', ''],
    ['周四', '概率论B 当前章节课后习题', '体育到课(14:45)\n晚间：码医1.5h+思政30min', '休息', '体育课已覆盖'],
    ['周五', '微积分B 课后习题+整理错题', '英语读写 写作1篇(200词+)', '乒乓球90min', ''],
    ['周六', '微积分B 本周错题二刷', '码医3h + 股票周复盘', '力量C 全身+敏捷', '晚间自由'],
    ['周日', '概率论B 本周错题二刷', '英语录音15min(面试练习)', '自由/补缺', '周复盘+小艺推送'],
])
P('', size=4)

# === 五 ===
H1('五、科目优先级与策略')
P('v6调整：思政提至P2（背了就有分），英语降至P3（CET-4已过），码医新增为P1。', size=10, color=GRAY)
P('', size=2)
TABLE(['优先级', '科目', '每周投入', '策略'], [
    ['P0 最重要', '微积分B', '8-10h', '转专业硬门槛80分+。每天必碰。B站宋浩+教材+真题。'],
    ['P0 次重要', '概率论与数理统计B', '5-6h', '逐章推进，错题标注，周末二刷。公式理解其意。'],
    ['P1', '微观经济学', '4-5h', '面试知识储备+股票基本面。塔勒布杠铃策略，学了不亏。'],
    ['P1', 'Python程序设计', '3-4h', '周二到课+完成作业。码医需要Python功底。'],
    ['P1 新', '码医(MediCode)', '5-6h', '面试核心素材。周二/四晚各1.5h+周六3h。'],
    ['P2 上调', '思法/近现代史', '2h/周', '从P3上调。考前突击收益极高。每周过一章。'],
    ['P3 下调', '英语读写/视听说', '1-2h', '从P2下调。CET-4已过，保持即可。'],
    ['P3', '形策/职业规划/选修', '考前1天', '不纳入周计划。考前集中突击。'],
])
P('', size=4)

# === 六 ===
H1('六、训练方案（精简版）')
P('所有力量训练围绕乒乓球发力链：腿 -> 腰 -> 肩 -> 臂 -> 腕。热身5min+训练+拉伸10min。', size=10, color=GRAY)

P('力量A -- 腿+核心（周一）', bold=True, size=11, color=BLUE)
TABLE(['动作', '组数x次数', '对乒乓球的作用'], [
    ['高脚杯深蹲', '4x12', '模拟准备姿势下蹲，练蹬转发力'],
    ['保加利亚分腿蹲', '3x10/腿', '单腿蹬地爆发力'],
    ['侧向箭步蹲', '3x10/侧', '核心动作--直接练台前左右移动'],
    ['哑铃罗马尼亚硬拉', '3x12', '臀腿后侧，降低重心稳定性'],
    ['平板+侧支撑', '各60s x3', '核心稳定，转腰躯干不散'],
    ['俄罗斯转体(持哑铃)', '3x20', '模拟正手拉球转腰发力'],
])
P('力量B -- 上肢+旋转（周三）', bold=True, size=11, color=BLUE)
TABLE(['动作', '组数x次数', '对乒乓球的作用'], [
    ['哑铃卧推', '3x10', '手臂推挡力量基础'],
    ['坐姿哑铃推举', '3x12', '肩部耐力，连续拉球不酸'],
    ['单臂哑铃划船', '3x12/侧', '反手位发力支撑'],
    ['弹力带肩外旋', '3x15', '必做！防肩伤关键动作'],
    ['腕弯举+反腕弯举', '3x15', '握拍和前臂精细控制'],
    ['药球侧抛(对墙)', '3x8/侧', '正手转腰爆发力模拟'],
])
P('力量C -- 全身+敏捷（周六）', bold=True, size=11, color=BLUE)
TABLE(['动作', '组数x次数', '对乒乓球的作用'], [
    ['壶铃摆举/哑铃高翻', '4x8', '全身蹬伸爆发力，腿->腰->臂一条链'],
    ['跳箱/原地跳蹲', '3x8', '启动第一步爆发，接发球那一下'],
    ['弹力带侧向滑步', '3x30s/向', '台前左右步法速度和耐力'],
    ['俯身哑铃飞鸟', '3x12', '肩胛骨稳定，防圆肩驼背'],
    ['农夫行走', '3x40s', '握力耐力+全身刚性'],
])
P('乒乓球技术（周二、周五，90min）', bold=True, size=11, color=BLUE)
P('热身10 -> 正手20 -> 反手20 -> 发球接发10 -> 实战/多球30。重点：蹬->转->收发力链，不要只用手臂打球。', size=10)
P('', size=4)

# === 七 ===
H1('七、转专业面试准备')
P('面试随时可能通知。每周日下午15分钟英语录音，不叫模拟面试。门槛越低，越可能真的做。', size=10, color=GRAY)
TABLE(['周次', '录音内容'], [
    ['第1周(5/29-6/4)', '英语自我介绍1-2min：我是谁、项目(缓伴/码医)、为什么转经济学国际投资方向'],
    ['第2周(6/5-6/11)', '项目深挖：码医DRG引擎怎么做？203份测试数据来源？94.1%准确率怎么测的？'],
    ['第3周(6/12-6/18)', '叙事线：实践撞墙->发现理论缺口->经济学是答案。物流到经济学连接能讲清吗？'],
])
P('', size=4)

# === 八 ===
H1('八、降载方案')
P('今天只做一道微积分题 + 15分钟任何运动 = 赢了', bold=True, size=12, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)
P('^ 贴床头。设成手机壁纸。^', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
P('', size=2)
P('每周末检查三项（3分钟）：', bold=True)
B('本周任务卡完成率 < 60%？')
B('睡眠质量 < 5/10？（连续入睡困难或早醒）')
B('情绪状态 < 5/10？（持续低落、莫名烦躁）')
P('-> 任意两项不达标 -> 下周降载：学习4h/天，训练3次/周，微积分至少碰一道题。')
P('-> 降载后仍撑不住 -> 再降。降到一道题+15分钟运动。照样是赢。', bold=True)
P('', size=4)

# === 九 ===
H1('九、手机物理隔离')
P('不是意志力问题。手机在床边就不可能早睡。只有一个规则：', size=10, color=GRAY)
P('22:30 手机移出卧室 -> 实体闹钟替代 -> 23:00 熄灯', bold=True, size=13, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
B('22:30-23:00 阅读/拉伸：纸质书或Kindle。教材也行，不看电子屏幕。')
B('23:00 强制关灯。睡不着躺着发呆也比刷手机好，至少大脑在休息。')
B('早起第一件事不是看手机：洗漱、喝水、完成晨间流程后再碰。')
P('', size=4)

# === 十 ===
H1('十、AI 助理（小艺 Claw）+ 手机闹钟双保险')
P('小艺是真实的AI助理，负责每日推送与追踪。手机闹钟作为离线备份。', size=10, color=GRAY)
P('', size=2)
TABLE(['时间', '小艺推送', '手机备份'], [
    ['每日 9:00', '推送当日任务卡确认', '手机闹钟 9:00，标签=今日任务'],
    ['每日 21:30', '推送完成度检查 + 心情收集', '手机闹钟 21:30 + 进度表打钩'],
    ['每日 22:30', '推送：手机移出卧室时间到了', '手机闹钟 22:30，响了就放客厅'],
    ['每周日 21:00', '推送周复盘提醒（三检 + 英语录音）', '手机日历每周重复事件'],
    ['降载预警', '连续2天未回复 -> 自动推送降载提醒', '自己看进度表，缺2天以上触发'],
])
P('', size=2)
P('手机闹钟是底线。小艺在线时双重确认，小艺不在时手机接棒。', size=9, color=GRAY)
P('', size=4)

# === 十一 ===
H1('十一、三周递进节奏')
TABLE(['', '第1周 5/29-6/4\n压力测试周', '第2周 6/5-6/11\n全速推进周', '第3周 6/12-6/18\n考前冲刺周'], [
    ['定位', '建立任务卡习惯。先做到每天都写两张卡。内容少没关系，关键是形成节奏。', '提高完成率和质量。开始回头复习微积分和概率论前面章节。', '刷真题+错题二刷。模拟考试状态限时做题。考前48h突击思政。'],
    ['学习', '跟上各科进度。微积分每天>=1h。不追求量，追求开始做。', '两张任务卡全完成。周末系统性二刷错题。弱项章节突破。', '真题限时训练。错题至少两遍。思政翻重点。'],
    ['训练', '轻重量，学动作，找发力感。乒乓球找手感，恢复节奏。', '加重量，每组做到力竭前1-2次停。乒乓球加反手衔接。', '维持重量，组数可缩减。乒乓球1-2次，实战为主。'],
    ['节点', '周日：第一次英语录音\n周日：压力测试复盘\n->决定第2周是否降载', '周日：第二次英语录音\n复查微积分弱项\n针对性补强', '周日：第三次英语录音\n保证睡眠>=8h/天\n考前状态调整'],
])
P('', size=4)

# === 十二 ===
H1('十二、每日进度追踪表')
P('每天睡前30秒填一行。一堆钩攒在一起，就是继续下去的动力。', size=10, color=GRAY)
dates = ['5/29周五','5/30周六','5/31周日','6/1周一','6/2周二','6/3周三','6/4周四',
         '6/5周五','6/6周六','6/7周日','6/8周一','6/9周二','6/10周三','6/11周四',
         '6/12周五','6/13周六','6/14周日','6/15周一','6/16周二','6/17周三','6/18周四']
TABLE(['日期', '任务卡1', '任务卡2', '训练', '一句话心情'],
      [[d,'','','',''] for d in dates])
P('', size=4)

# === FOOTER ===
DIVIDER()
P('董事会：大野耐一（流程效率）· 塔勒布（尾部风险）· 张小龙（人性设计）· 小艺 Claw（AI助理与执行追踪）', size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
P('v6 合并优化：融合原版v3与改进版v5 -> 修三周递进表结构 + 精简训练方案 + 手机闹钟双保险 + 小艺确认保留', size=8, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# === SAVE ===
output = os.path.join(os.path.dirname(__file__), '..', 'output', '三周综合计划_v6.docx')
os.makedirs(os.path.dirname(output), exist_ok=True)
doc.save(output)
print('Saved: ' + output)
