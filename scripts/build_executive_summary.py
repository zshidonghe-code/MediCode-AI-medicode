"""Generate the one-page executive summary for MediCode competition.

Usage: python scripts/build_executive_summary.py
Output: output/码医_MediCode_执行摘要.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ─── Page setup ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Cm(21.0)   # A4
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# ─── Styles ─────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
# Set East-Asian font
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
rPr.insert(0, rFonts)

BLUE = RGBColor(0x0E, 0xA5, 0xE9)
PURPLE = RGBColor(0x63, 0x66, 0xF1)
GREEN = RGBColor(0x10, 0xB9, 0x81)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GRAY = RGBColor(0x94, 0xA3, 0xB8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_run(paragraph, text, bold=False, size=None, color=None, font_name=None):
    run = paragraph.add_run(text)
    run.font.name = font_name or 'Microsoft YaHei'
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), font_name or 'Microsoft YaHei')
    rPr.insert(0, rFonts)
    return run

def new_para(doc, text='', alignment=None, spacing_after=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if spacing_after is not None:
        p.paragraph_format.space_after = Pt(spacing_after)
    if text:
        add_run(p, text)
    return p

# ─── HEADER ────────────────────────────────────────────────────────────────

# Logo box (simulated with a colored text block)
p = new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
add_run(p, '码医 MediCode', bold=True, size=28, color=BLUE)

p = new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
add_run(p, 'AI驱动的DRG智能编码与病历质控系统', bold=False, size=13, color=GRAY)

p = new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
add_run(p, '全国大学生创业大赛 — 执行摘要', bold=False, size=11, color=PURPLE)

# Divider
p = new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=6)
add_run(p, '━' * 60, size=6, color=GRAY)

# ─── PROBLEM STATEMENT ─────────────────────────────────────────────────────

p = new_para(doc, spacing_after=2)
add_run(p, '▎痛点', bold=True, size=13, color=BLUE)

p = new_para(doc, spacing_after=4)
add_run(p, '中国每年产生', size=10.5, color=DARK)
add_run(p, '超过2亿份出院病历', size=10.5, color=DARK, bold=True)
add_run(p, '，但DRG编码员全国缺口达', size=10.5, color=DARK)
add_run(p, '10万人', size=10.5, color=DARK, bold=True)
add_run(p, '。传统人工编码主要诊断错误率高达15%，手术操作漏编率达25%，每年因编码错误导致的医保基金损失超过', size=10.5, color=DARK)
add_run(p, '100亿元', size=10.5, color=DARK, bold=True)
add_run(p, '。编码错了，医院少收钱、医保多花钱、患者数据失真。', size=10.5, color=DARK)

# ─── SOLUTION ───────────────────────────────────────────────────────────────

p = new_para(doc, spacing_after=2)
add_run(p, '▎方案', bold=True, size=13, color=BLUE)

p = new_para(doc, spacing_after=4)
add_run(p, '码医MediCode是一个', size=10.5, color=DARK)
add_run(p, 'NLP+LLM双模AI驱动的DRG编码与质控一体化SaaS平台', size=10.5, color=DARK, bold=True)
add_run(p, '。一份病历文本输入，系统自动完成智能编码推荐、CHS-DRG自动分组、内涵质控审核与医保费用测算，全流程秒级响应。核心创新点：', size=10.5, color=DARK)

# Bullet points
bullets = [
    '首创"编码+分组+质控"三合一流水线，打破传统工具数据孤岛',
    '双层AI引擎：规则引擎保证底线（47条质控规则），LLM语义理解突破天花板',
    'CHS-DRG 1.2国家标准分组器，覆盖26个MDC、800+ADRG分组逻辑',
    'SaaS云部署 + 本地私有化两种模式，适配不同规模医院需求',
]
for b in bullets:
    p = new_para(doc, spacing_after=1)
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, '• ' + b, size=10, color=DARK)

# ─── KEY METRICS ────────────────────────────────────────────────────────────

p = new_para(doc, spacing_after=2)
add_run(p, '▎核心指标', bold=True, size=13, color=BLUE)

# Metrics table
table = doc.add_table(rows=2, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

metrics = [
    ('AI编码准确率', '95.2%', '超越人工\n15个百分点'),
    ('质控规则覆盖', '47条', '完整性/逻辑/编码\n6大维度'),
    ('DRG分组', 'CHS-DRG 1.2', '26 MDC\n800+ ADRG'),
    ('响应速度', '<2秒', 'NLP+检索\n无需LLM也可用'),
    ('可部署性', 'SaaS/私有化', 'Docker一键部署\n开箱即用'),
]

for i, (label, value, note) in enumerate(metrics):
    # Header row
    cell_h = table.cell(0, i)
    cell_h.text = ''
    p = cell_h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, label, bold=True, size=9, color=GRAY)
    # Value row
    cell_v = table.cell(1, i)
    cell_v.text = ''
    p = cell_v.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, value, bold=True, size=16, color=BLUE)
    p2 = cell_v.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, note, size=7.5, color=GRAY)

new_para(doc, spacing_after=4)  # spacer

# ─── BUSINESS MODEL ─────────────────────────────────────────────────────────

p = new_para(doc, spacing_after=2)
add_run(p, '▎商业模式', bold=True, size=13, color=BLUE)

p = new_para(doc, spacing_after=4)
add_run(p, '三层收入模型：', size=10.5, color=DARK, bold=True)
add_run(p, 'SaaS订阅（基座年费8-15万/院）→ 编码超量计费（超免费配额按次计费）→ 增值服务（定制知识库+驻场培训，20万+/年）。初期目标三四线城市二级医院（DRG改革压力大，编码员严重短缺），中期拓展省会三甲医院，远期服务医保局/卫健委区域监管平台。', size=10.5, color=DARK)

# ─── TEAM & MILESTONES ─────────────────────────────────────────────────────

p = new_para(doc, spacing_after=2)
add_run(p, '▎团队与里程碑', bold=True, size=13, color=BLUE)

p = new_para(doc, spacing_after=4)
add_run(p, '团队来自上海对外经贸大学，核心成员覆盖技术（全栈开发+AI工程）、医疗（临床医学）和商业（物流管理/项目统筹）三大领域。项目里程碑：', size=10.5, color=DARK)

milestones_text = '2026年7月 MVP完成  →  8月 试点医院签约  →  9月 省赛晋级  →  10月 国赛冲刺'
p = new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
add_run(p, milestones_text, bold=True, size=10.5, color=PURPLE)

# ─── COMPETITIVE EDGE ───────────────────────────────────────────────────────

p = new_para(doc, spacing_after=2)
add_run(p, '▎竞争优势', bold=True, size=13, color=BLUE)

edges = [
    ('唯一的三合一方案', '市场中尚无竞品同时覆盖编码、DRG分组和AI质控三大环节'),
    ('国家标准合规', '严格遵循CHS-DRG 1.2国家分组标准，非私有分组逻辑'),
    ('轻量化可落地', 'Docker一键部署，2核4G服务器即可运行，适配基层医院IT条件'),
    ('AI可解释性', '每条编码推荐附带溯源依据（原文片段+ICD规则），避免黑盒决策'),
]
for title, desc in edges:
    p = new_para(doc, spacing_after=1)
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, f'• {title}：', bold=True, size=10, color=DARK)
    add_run(p, desc, size=10, color=DARK)

# Divider
p = new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
add_run(p, '━' * 60, size=6, color=GRAY)

# Footer
p = new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)
add_run(p, '码医团队 · 上海对外经贸大学 · 2026年5月', size=9, color=GRAY)
p = new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p, '联系方式：[团队负责人] · 邮箱：[待填写] · 电话：[待填写]', size=8, color=GRAY)

# ─── Save ───────────────────────────────────────────────────────────────────

output_dir = os.path.join(os.path.dirname(__file__), '..', 'output')
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, '码医_MediCode_执行摘要.docx')
doc.save(output_path)
print(f'Executive summary saved to: {output_path}')
