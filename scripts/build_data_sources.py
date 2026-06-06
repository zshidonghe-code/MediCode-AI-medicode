# coding: utf-8
"""Generate standalone data source traceability table for Q&A defense."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

sty = doc.styles['Normal']
sty.font.name = '微软雅黑'; sty.font.size = Pt(10.5)
sty.paragraph_format.space_after = Pt(4); sty.paragraph_format.line_spacing = 1.2
sty.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

BLUE = RGBColor(0x0E,0xA5,0xE9); DARK = RGBColor(0x1E,0x29,0x3B)
GRAY = RGBColor(0x94,0xA3,0xB8); GREEN = RGBColor(0x10,0xB9,0x81)

def P(text, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    return p

P('码医 MediCode', bold=True, size=24, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
P('路演答辩 · 数据溯源表', size=14, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
P('评委问数据来源时，直接翻到对应行回答。不需要记住所有数字，知道在哪就行。', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
P('', size=6)

# Category 1: 行业规模
P('一、行业规模与政策依据', bold=True, size=13, color=BLUE)
P('', size=2)

t1 = doc.add_table(rows=7, cols=4)
t1.style = 'Light Grid Accent 1'; t1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['路演中出现的数字', '数据来源', '发布年份', '评委追问时的应答']):
    c = t1.rows[0].cells[i]; c.text = ''
    r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

src_data = [
    ['全国10万编码员缺口', '中国医院协会病案管理专业委员会\n《全国病案科人力资源调查报告》', '2024', '行业共识数据。三甲医院平均3-5名编码员，年出院15万人次，合理配置需8-10人。'],
    ['人工编码主诊断错误率10-15%', '《中国病案》期刊多篇研究\n国家医保局飞行检查通报', '2020-2024', '学界主流区间。实际飞检中发现的编码错误率更高，10-15%是保守估计。'],
    ['手术操作漏编率20-30%', '多地医保飞行检查通报\n《中国医疗保险》期刊', '2023-2024', '次要手术操作极易漏编。一次飞检倒查3年病历，漏编的直接后果是医保拒付。'],
    ['每年编码错误损失超100亿', '国家医保局基金监管年报\n行业测算（基于3万亿总支出×编码错误率×平均支付偏差）', '2024', '保守估算。3万亿×10%错误率×3.5%支付偏差=105亿。实际可能更高。'],
    ['DRG/DIP 2025年底全面覆盖', '国家医保局《DRG/DIP支付方式改革\n三年行动计划》（医保发〔2021〕48号）', '2021', '政策原文：2025年底前，所有符合条件的住院医疗机构全面推行DRG/DIP付费。'],
    ['全国二级以上医院15,000+家', '国家卫健委《全国医疗卫生机构\n统计年报》', '2024', '含二级医院约11,000家+三级医院约3,500家。不含一级医院和社区卫生中心。'],
]
for ri, row in enumerate(src_data):
    for ci, val in enumerate(row):
        c = t1.rows[ri+1].cells[ci]; c.text = ''
        r = c.paragraphs[0].add_run(val); r.font.size = Pt(8.5)
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

P('', size=4)

# Category 2: 技术与产品
P('二、技术与产品数据', bold=True, size=13, color=BLUE)
P('', size=2)

t2 = doc.add_table(rows=6, cols=4)
t2.style = 'Light Grid Accent 1'; t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['路演中出现的数字', '数据来源', '测试条件', '评委追问时的应答']):
    c = t2.rows[0].cells[i]; c.text = ''
    r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

tech_data = [
    ['诊断Top-1准确率94.1%', '码医内部基准测试\n203份模拟临床病历', '8科室(心内/呼吸/消化/内分泌\n/神内/骨科/普外/妇产科)\nLLM: Qwen2.5-7B本地部署\n规则引擎兜底+LLM增强双模式', '203份病历的编码结果和评分表可提供。\n测试不是"挑好的报"，覆盖了8个常见科室。\n消化内科和普外科存在少量混淆，已记录在案。'],
    ['诊断F1得分84.5%', '同上', 'Top-3推荐下的综合得分\n精确率+召回率加权', 'F1比Top-1更能反映综合表现。\n84.5%意味着AI推荐的Top-3编码中，\n绝大多数情况下包含了正确答案。'],
    ['手术编码F1=23.5%', '同上，手术编码部分', '571条高频手术编码\n覆盖9个专科', '坦诚公开。手术编码库庞大(3000+条)，\n目前优先覆盖高频编码。不依赖LLM也可\n通过规则引擎保底。省赛前目标提至40%+。'],
    ['47条质控规则，6大维度', '码医质控引擎设计文档', '完整性12+逻辑一致性8+编码一致性7\n+时效性5+规范表达9+语义质量6', '规则来自国家病案质控标准+临床指南。\n每条规则都有明确的触发条件和缺陷等级。'],
    ['ICD-10仅覆盖920条(3-5%)', 'ICD-10全量编码集约2-3万条\n码医当前覆盖920条', '920条=前80%高频诊断\n覆盖95%+出院病例\n低频罕见病：LLM语义兜底', '诚实承认：覆盖面确实不够。但策略明确——\n先吃掉80%最常见诊断(已做到)，低频罕见病\n用LLM推理兜底(大模型见过全量ICD知识)。\n编码库设计为模块化JSON，每月迭代扩展。\n这不是"做不了"，是"在排队"。'],
]
for ri, row in enumerate(tech_data):
    for ci, val in enumerate(row):
        c = t2.rows[ri+1].cells[ci]; c.text = ''
        r = c.paragraphs[0].add_run(val); r.font.size = Pt(8.5)
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

P('', size=4)

# Category 3: 竞品价格
P('三、竞品价格参考', bold=True, size=13, color=BLUE)
P('', size=2)

t3 = doc.add_table(rows=5, cols=3)
t3.style = 'Light Grid Accent 1'; t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['竞品', '价格区间', '来源']):
    c = t3.rows[0].cells[i]; c.text = ''
    r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

comp_data = [
    ['东软望海 DRG模块', '50-200万/年（含HIS捆绑）', '行业公开信息+医院信息化招标公告'],
    ['国新健康 DRG数据服务', '30-100万/年', '上市公司年报+公开报价'],
    ['森亿智能 AI病历', '50-150万/年', '行业媒体报道+融资披露信息'],
    ['小型编码SaaS工具', '2-10万/年', '各厂商官网公开定价'],
]
for ri, row in enumerate(comp_data):
    for ci, val in enumerate(row):
        c = t3.rows[ri+1].cells[ci]; c.text = ''
        r = c.paragraphs[0].add_run(val); r.font.size = Pt(8.5)
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

P('', size=8)
P('━'*50, size=6, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
P('使用方式：打印一份带进答辩现场。评委问"这个数字哪来的"，翻到对应行，念出来就行。', size=10, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
P('码医团队 · 2026年5月', size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

output = os.path.join(os.path.dirname(__file__), '..', 'output', '码医_答辩数据溯源表_v2.docx')
os.makedirs(os.path.dirname(output), exist_ok=True)
doc.save(output)
print('Saved: ' + output)
