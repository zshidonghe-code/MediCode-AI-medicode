# -*- coding: utf-8 -*-
"""Generate invitation letter v2 for Jin Hua - MediCode medical advisor
v2 changes (per council verdict 2026-06-01):
  1. Frame flip: from "we need you" to "your expertise → AI + paper"
  2. Value reframe: from "only 5 hours" to "国赛署名+论文+课题"
  3. Social proof placeholder: [经XX引荐]
  4. Call to action: from "期待回复" to specific meeting request
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

LQ = '“'  # Chinese left double quote
RQ = '”'  # Chinese right double quote

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), 'SimSun')
rPr.insert(0, rFonts)


def add_title(text, size=22, color=0x1a3c6e):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor((color >> 16) & 0xff, (color >> 8) & 0xff, color & 0xff)
    return p


def add_heading_text(text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    return p


def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_list_item(num, text):
    p = doc.add_paragraph()
    run = p.add_run(f'  {num}. {text}')
    run.font.size = Pt(11)
    return p


# =====================================================
# CONTENT (v2 — council-advised rewrite)
# =====================================================

add_title('邀 请 函')

greeting = doc.add_paragraph()
greeting.paragraph_format.space_before = Pt(12)
run = greeting.add_run('尊敬的金花老师：')
run.font.size = Pt(12)
run.font.bold = True

# ---- OPENING: frame flip ----
# Old: "我是学生，我需要你的帮助"
# New: "你的经验，可以变成AI系统的一部分"
opening = (
    '您好！[经XX引荐]，冒昧致信。我是上海对外经贸大学的大一学生郑诗东和，'
    '正在做一个与您专业方向高度契合的AI医疗项目，想邀请您共同参与。'
)
add_body(opening)

# Core pitch — Flip the frame entirely
pitch = (
    '您在上海市全科医学临床质控中心担任秘书，每天的工作是制定质控标准、检查病历质量、'
    '推动多病共存诊疗决策的规范化。而我们团队开发的' + LQ + '码医 MediCode' + RQ + '系统，'
    '用人工智能自动完成ICD-10诊断编码、DRG付费分组和病历内涵质控——本质上，'
    '就是把像您这样的全科质控专家的判断逻辑，变成可复用的AI规则引擎。\n\n'
    '目前系统已集成920条诊断编码和611条手术编码，基于CHS-DRG 1.2方案实现了'
    '从NLP语义解析到DRG分组的全流程自动化，编码Top-1准确率94.1%。'
    '但AI能做的是' + LQ + '算得快' + RQ + '，它缺的是' + LQ + '判断得准' + RQ + '——'
    '而这恰恰是您每天都在做的事。'
)
add_body(pitch)

# ---- WHY HER (keep the table, strengthen language) ----
add_body(
    '在寻找合作专家的过程中，我发现您的专业积累与这个项目的技术方向有着罕见的精准匹配：'
)

table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table.autofit = True

hdr = table.rows[0].cells
hdr[0].text = '您的专业积累'
hdr[1].text = '如何直接转化为AI能力'
for cell in hdr:
    for pp in cell.paragraphs:
        for r in pp.runs:
            r.font.bold = True
            r.font.size = Pt(10)

data = [
    ('上海市全科医学临床质控中心秘书',
     '您制定的质控标准 → 转化为AI质控引擎的规则逻辑，让机器学会' + LQ + '什么叫一份好病历' + RQ),
    ('多病共存诊疗决策研究',
     '合并症/并发症（CC/MCC）识别是全科诊疗和DRG编码的共同难点——您的临床判断可以直接校准AI编码逻辑'),
    ('同济大学附属医院 · 博导 · 副教授',
     '学术训练意味着您理解什么样的产出能发表、能申报课题——这正是我们需要的合作深度'),
    ('上海本地 · 杨浦区',
     '同城意味着可以真正协作，而不是' + LQ + '挂个名就结束' + RQ + '的纸面关系'),
]
for i, (left, right) in enumerate(data):
    row = table.rows[i + 1]
    row.cells[0].text = left
    row.cells[1].text = right
    for cell in row.cells:
        for pp in cell.paragraphs:
            for r in pp.runs:
                r.font.size = Pt(10)

doc.add_paragraph()

# ---- COMPETITION AS AMPLIFIER (not as the ask) ----
add_body(
    '我们正在备战中国国际大学生创新大赛（原' + LQ + '互联网+' + RQ + '），目标是全国第一名。'
    '为什么这对您也重要？因为国赛获奖项目的指导老师署名，本身就是职称评审和学术影响力的加分项。'
    '更重要的是——一个拿了国赛奖的AI医疗项目，可以自然延伸出论文发表和课题申报的合作空间。'
)

# ---- WHAT WE'RE PROPOSING (reframed as collaboration, not favor) ----
add_heading_text('我们提议的合作方式：')

duties = [
    '指导老师署名：在竞赛报名材料中作为指导老师/医学顾问署名——这是最直接的学术影响力体现；',
    '临床逻辑审核（1-2次）：审阅AI编码推荐和质控规则的医学合理性——您的质控经验变成AI规则的第一道质量把关；',
    '方向性指导（可选）：如时间允许，在路演答辩前给予方向性建议——您的表达本身就是最好的答辩培训。',
]
for i, d in enumerate(duties, 1):
    add_list_item(i, d)

add_body(
    '整个合作周期（2026年6月-10月）在时间投入上非常克制，但在学术产出上我们希望做到扎实。'
    '合作的核心不是' + LQ + '请您帮忙' + RQ + '，而是' + LQ + '您的质控经验值得被转化为AI，'
    '而我们可以成为这个转化的技术方' + RQ + '。'
)

# ---- VALUE FOR HER (completely rewritten per council) ----
add_heading_text('具体而言，这次合作对您的学术价值：')

values = [
    '国赛获奖署名：中国国际大学生创新大赛国赛获奖项目的正式指导老师——直接计入学术履历和职称评审材料；',
    '交叉学科论文：'
    + LQ + '全科质控规则的AI化方法' + RQ + '本身就是一个高质量的研究课题——从临床经验到算法规则的转化过程，'
    '可以产出医学信息学/全科医学方向的论文，您作为通讯作者；',
    '课题预研基础：全科质控AI化可以成为后续国自然青年/面上项目或省部级课题的预研基础——'
    '我们在技术上已经把路趟了一遍，您在临床上的把关让这条路变得可信；',
    '社会影响力：您的名字出现在一个服务医院的AI系统里——质控专家的专业判断通过AI放大到全院、全行业。',
]
for i, v in enumerate(values, 1):
    add_list_item(i, v)

# ---- ATTACHMENTS & NEXT STEP ----
add_body(
    '随信附上项目的商业计划书、执行摘要和系统Demo录屏。'
    '文字和截图不如当面交流直观——如果您本周四或下周有时间，'
    '我可以带着电脑到杨浦区中心医院，用15分钟当面演示系统并听取您的判断。'
)

# ---- CLOSING ----
closing = doc.add_paragraph()
closing.paragraph_format.space_before = Pt(12)
run = closing.add_run(
    '无论您最终是否方便参与，都非常感谢您花时间了解我们的项目。\n'
    '一个全科质控专家的名字出现在AI医疗项目的指导老师栏里——'
    '这不仅是我们团队的荣幸，也将是评审眼中' + LQ + '医学+AI跨界融合' + RQ + '的最佳注脚。\n\n'
    '期待能与您当面交流。'
)
run.font.size = Pt(11)

# ---- SIGNATURE ----
sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sig.paragraph_format.space_before = Pt(24)
today = datetime.date.today().strftime('%Y年%m月%d日')
run = sig.add_run(
    '郑诗东和\n'
    '上海对外经贸大学 · 物流管理（中澳合办）· 大一\n'
    '码医 MediCode 项目负责人\n'
    '手机：15800565959\n'
    '邮箱：1975790036@qq.com\n'
    f'日期：{today}'
)
run.font.size = Pt(10.5)

# Save
output_path = r'C:\Users\Donghe\Desktop\邀请函-金花-码医医学顾问.docx'
doc.save(output_path)
print(f'Done: {output_path}')
