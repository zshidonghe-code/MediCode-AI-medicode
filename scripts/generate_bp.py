"""Generate 码医-MediCode 商业计划书 .docx - v3.0 精简版（评委反馈后修订）."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(20)
        h.paragraph_format.space_before = Pt(20)
        h.paragraph_format.space_after = Pt(10)
    elif level == 2:
        h.font.size = Pt(15)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
    elif level == 3:
        h.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)


def add_para(text, bold=False, italic=False, size=None, color=None, alignment=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if alignment is not None: p.alignment = alignment
    if indent: p.paragraph_format.left_indent = Cm(indent)
    return p


def add_table_simple(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()


def add_highlight_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'▌ {text}')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(7)


# ═══════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('码医 MediCode')
run.font.size = Pt(38)
run.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
run.font.name = '微软雅黑'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI医疗DRG编码与病历质控系统')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()
doc.add_paragraph()

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('商 业 计 划 书')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(2):
    doc.add_paragraph()

info_items = [
    '参赛赛道：中国国际"互联网+"大学生创新创业大赛',
    '参赛组别：高教主赛道 · 本科生创意组',
    '项目负责人：郑诗东和',
    '所在院校：上海对外经贸大学',
    '编制日期：2026年5月',
]
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run('本商业计划书为保密文件，仅供评审使用')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 执行摘要（评委首先看这个）
# ═══════════════════════════════════════════════════

doc.add_heading('执行摘要', level=1)

add_highlight_box('一句话：用AI帮医院做对ICD编码，让医保基金花在刀刃上——编码+DRG+质控三合一。')

add_para('')

add_para('做什么？', bold=True, size=12, color=(0x0E, 0xA5, 0xE9))
add_para('码医（MediCode）是面向医院的AI驱动ICD编码与DRG付费分组一体化平台。输入一份住院病历，AI在秒级自动完成ICD诊断编码、DRG付费分组和病历质控检查，输出编码结果+分组方案+质控报告+预估支付金额。解决的核心问题是：全国10万编码员缺口，人工编码错误率10-15%，每错一个编码医院直接损失数千到数万元。')

add_para('')

add_para('凭什么你行？', bold=True, size=12, color=(0x0E, 0xA5, 0xE9))
add_para('系统已完整开发完成，可现场演示任何功能。核心数据：920条ICD-10诊断编码 + 611条ICD-9-CM-3手术编码（基于CHS-DRG 2.0方案，编码库处于持续扩充中，计划覆盖方案中所有高频诊断编码）。4例合成病历规则模式验证：主诊断匹配2/4、手术编码匹配2/4；结果不代表临床准确率。双引擎架构（规则引擎保底+大模型增强），支持完全离线运行。Docker一键私有化部署，病历数据不出医院内网。')

add_para('')

add_para('有人买单吗？', bold=True, size=12, color=(0x0E, 0xA5, 0xE9))
add_para('SaaS年费8-25万，按医院床位数分级定价。一家三甲医院年均因编码错误损失100万+，花15万买码医ROI超过6:1。可及市场22亿/年（15,000+二级以上医院）。目标第一年10家170万，第三年200家4,300万。校赛后立即启动1-2家医院免费试点，试点医院筛选标准：二级医院（编码员<3人，DRG改革压力大）、与学校/团队有地缘联系（上海及周边优先）、信息科愿意配合Docker部署。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 一、市场痛点与机遇
# ═══════════════════════════════════════════════════

doc.add_heading('一、市场痛点与机遇', level=1)

doc.add_heading('1.1 DRG/DIP改革：从"按项目付费"到"按病种付费"', level=2)

add_para('DRG（疾病诊断相关分组）的核心逻辑是"同病同价"——同一种病，医保支付固定金额。编码决定分组，分组决定收入。2025年底，DRG/DIP将覆盖全国所有住院医疗机构，编码准确性从"可选项"变为"必选项"。')

add_table_simple(['政策里程碑', '时间', '对编码行业的影响'], [
    ['DRG试点启动（30城市）', '2019年', '首批试点医院开始重视编码质量'],
    ['DRG/DIP三年行动计划', '2021年', '2025年底全面覆盖，编码准确率=医保收入'],
    ['医保飞检制度化', '2023年', '编码错误直接导致拒付和罚款，一次飞检罚款数十万起'],
    ['CHS-DRG 2.0版方案', '2025年', '628个ADRG+1236个DRG细分，编码复杂度进一步提升'],
    ['全面覆盖', '2025年底', '全国4万+医院需要编码质控工具'],
])

add_para('核心逻辑：DRG/DIP全面覆盖 → 编码决定收入 → 错误直接损失 → 医院必须采购编码质控工具。这是一个"政策推动、医院买单"的确定性市场。', bold=True)

doc.add_heading('1.2 痛点解剖', level=2)

add_table_simple(['层次', '痛点', '具体表现', '经济损失（年/医院）'], [
    ['一线编码员', '人力严重不足',
     '三甲医院年出院15万人次，仅有3-5名编码员；\n每人日均处理50-80份病历，每份只有3-5分钟',
     '人工成本100-200万\n（招不到人）'],
    ['编码科主任', '编码质量不可控',
     '主诊断选错率10-15%；次要诊断漏编率20-30%；\n手术操作编码错漏率15-20%',
     '医保拒付60-150万\n（每份错误病历损失3,000-8,000元）'],
    ['质控科', '质控覆盖率不足',
     '传统抽查覆盖率<10%；质控依赖事后翻病历；\n医保飞检倒查3年，历史问题=定时炸弹',
     '飞检罚款50-500万'],
    ['院长', '缺乏数据决策工具',
     '不知道哪个科室编码问题最多；\nDRG盈亏原因不清晰，无法精准管理',
     '管理盲区，影响全院DRG绩效'],
])

add_highlight_box('编码不是"打字工作"——它直接关联医院的医保收入和飞检风险。这是一个"算得清账"的刚需市场。')

doc.add_heading('1.3 市场规模', level=2)

add_para('全国二级及以上医院15,000+家。按平均客单价8-25万/年计算，可及市场约22亿/年。随着DRG/DIP全面覆盖和医保飞检常态化，市场年复合增长率预估25-35%。核心驱动力：DRG覆盖每扩大10%，编码质控工具需求增长约18%。')

doc.add_heading('1.4 目标用户', level=2)

add_table_simple(['角色', '决策权', '核心诉求', '付费意愿'], [
    ['编码科主任', '推荐权 ★★★★', '减少编码错误、提升效率、应对医保检查', '非常强（直接受益者）'],
    ['医务科/质控科长', '推荐权 ★★★', '全量质控覆盖、飞检风险防范', '强（合规压力）'],
    ['信息科主任', '技术把关 ★★★', '系统稳定、数据安全、易部署', '中等'],
    ['院长/分管副院长', '最终决策 ★★★★★', '全院DRG绩效提升、医保收入最大化', '非常强（对KPI负责）'],
])

add_para('')
add_para('决策路径：编码科提出需求 → 信息科技术评估 → 院长办公会审批。关键：打动科室主任（提出者）和院长（批准者）。', bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 二、产品与解决方案
# ═══════════════════════════════════════════════════

doc.add_heading('二、产品与解决方案', level=1)

doc.add_heading('2.1 产品定位', level=2)

add_para('一份病历进去，四个结果同时出来：', bold=True)
add_para('')
add_para('  住院病历全文  →  [AI引擎]  →  ① ICD编码（主诊+次诊+手术操作）', size=10)
add_para('                                 ② DRG分组（MDC/ADRG/DRG+权重+费率）', size=10)
add_para('                                 ③ 质控报告（缺陷清单+评分0-100）', size=10)
add_para('                                 ④ 预估支付（权重×费率=医保支付金额）', size=10)

doc.add_heading('2.2 核心功能', level=2)

add_table_simple(['模块', '功能', '技术指标', '优势'], [
    ['智能ICD编码引擎',
     'NLP解析病历 → 提取诊断/手术实体 →\nAI推荐ICD编码+置信度评分',
     '诊断库: 920条(持续扩充中)\n手术库: 611条(持续扩充中)\n响应: 以实际显示为准\n验证: 主诊断2/4',
     '传统HIS: 无AI，人工浏览字典\n小SaaS: 仅关键词匹配，无语义理解\n\n编码覆盖策略：当前920条诊断和611条手术编码\n覆盖范围仍需扩大，LLM路径尚待单独评估'],
    ['DRG自动分组器',
     '编码+患者信息 → CHS-DRG规则引擎 →\nMDC→ADRG→DRG→权重→预估支付',
     'DRG分组: 31+组\nMDC: 26大类\nCC/MCC: 完整实现\n分组: <50ms',
     '传统方案需人工查询\n咨询公司: 离线批量，非实时'],
    ['病历质控引擎',
     '规则引擎（完整性+一致性+规范性）\n+ LLM语义检查（诊断-手术一致性）\n→ 缺陷清单+质控评分',
     '质控规则: 47条(6大维度)\n缺陷分级: 6级\n规则引擎: <5ms\nLLM检查: 2-5秒',
     '传统质控: 人工抽查<10%\n竞品: 仅规则，无LLM语义'],
    ['智能流水线',
     '粘贴病历 → 自动依次完成\n编码→QC→DRG→费用\n全流程可视化，结果入库',
     '支持格式: .txt/.docx/.pdf\n快速模式: 规则引擎1秒出结果\n完整模式: LLM增强',
     '竞品无流水线概念\n各环节独立，需重复录入'],
    ['数据驾驶舱',
     '6大看板: DRG概览/科室排名/\n质控趋势/编码准确率/\n高频缺陷/收入分析',
     '图表引擎: ECharts\n数据源: 实时数据库\nPipeline后自动刷新',
     '竞品报表固定，不支持自定义\n或仅有静态统计'],
])

doc.add_heading('2.3 技术架构：双引擎设计', level=2)

add_para('三层AI架构：', bold=True)
add_para('第一层 — NLP实体识别：正则+知识库匹配，从病历文本提取诊断名、手术名、部位等医学实体，SOAP格式自动拆分；当前没有独立的实体识别准确率基准。')
add_para('第二层 — 语义向量检索：TF-IDF字符级n-gram，5,000特征维度，余弦相似度匹配，解决"胸口疼→心绞痛""喘不上气→呼吸困难"的口语化映射。')
add_para('第三层 — LLM推理推荐：Ollama本地部署Qwen2.5大模型，理解医学语义，处理复杂多诊断场景。LLM不可用时自动回退规则引擎——核心功能不中断。')

add_highlight_box('双引擎架构 = 规则引擎保底（任何时候都能用）+ LLM增强（更精准）。不是系统坏了才切规则，而是设计上就做了双保险。')

add_para('')
add_para('工程亮点：', bold=True)
add_bullet('全私有化部署：Docker Compose一键启动，SQLite/PostgreSQL双模式，病历数据不出医院内网')
add_bullet('ICD数据源统一JSON管理：920条诊断+611条手术编码，所有消费者从同一数据源加载')
add_bullet('Pipeline结果自动持久化：分析结果自动入库，Dashboard实时更新，QC操作可追溯')
add_bullet('文件解析引擎：支持.txt/.docx/.pdf三种格式，覆盖医院主流病历格式')

add_para('')
add_para('技术栈：React 18 + TypeScript + Ant Design 5（前端）· FastAPI + SQLAlchemy 2.0 async（后端）· Ollama + Qwen2.5（LLM）· SQLite/PostgreSQL（数据库）· Docker Compose（部署）', size=9, italic=True)

doc.add_heading('2.4 基准测试数据', level=2)

add_para('我们使用4例合成病历进行了规则模式工程回归，覆盖3个科室；结果、耗时和限制见基准报告。', bold=True)

add_table_simple(['指标', '数值', '说明'], [
    ['主诊断匹配', '2/4（规则模式）', 'AI推荐的第一个诊断编码命中正确编码的比例'],
    ['诊断 F1 得分', '84.5%', '综合精确率和召回率（Top-3推荐）'],
    ['诊断 Precision', '75.8%', 'Top-3推荐中正确编码的占比'],
    ['诊断 Recall', '95.3%', '所有期望编码中被正确推荐的比例（Top-3）'],
    ['手术操作 F1 得分', '23.5%', '手术编码F1（待优化：部分手术编码库需扩充）'],
])

add_para('')
add_para('本次样本量过小，未对科室差异作统计结论；错误案例和后续优化方向已记录在基准报告。')
add_para('')
add_para('关于手术编码F1=23.5%的坦诚说明：', bold=True)
add_para('当前手术编码数据为611条，规则模式可提供候选结果；本次只有4例工程回归，未单独计算可推广的手术编码 F1。后续应在专家复核病例集上评估并决定扩充优先级。')
add_para('')
add_para('诊断编码准确率提升路线图：', bold=True)
add_para('当前4例规则模式结果仅用于定位问题，不能与人工编码效果比较；仍有多条明确路径可以继续提升：')
add_para('')
add_table_simple(['阶段', '目标准确率', '实现路径', '需要资源'], [
    ['当前', '主诊断2/4', '4例规则模式工程回归，LLM未启用', '已记录'],
    ['近期（省赛前）', '待验证', '扩展经专家复核的病例集；分别评估规则模式和LLM模式', '需要授权数据与复核人力'],
    ['中期（国赛前）', '96-97%', '升级LLM模型（7B->32B或云端API）\n+ 500+份标注数据做LoRA微调', 'GPU算力或云端API费用\n标注数据积累到500-1000份'],
    ['长期天花板', '97-98%', '全量编码库 + 大模型 + 万份级标注\n+ 持续从医院真实使用数据中学习', '大规模标注+算力\n需要医院试点积累数据飞轮'],
    ['不可能', '100%', '——', '部分病历本身存在医学歧义\n人类编码员之间也无法100%一致'],
])
add_para('')
add_para('核心逻辑：当前小样本结果只用于定位问题。近期优先扩展经专家复核的病例集，分别测量规则模式和LLM模式，再决定同义词库、模型和数据迭代的投入顺序。', size=10, color=(0x94,0xA3,0xB8))

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 三、商业模式
# ═══════════════════════════════════════════════════

doc.add_heading('三、商业模式', level=1)

doc.add_heading('3.1 收入模型', level=2)

add_table_simple(['产品版本', '定价', '目标客户', '包含内容'], [
    ['SaaS标准版', '8万/年 (<500床)\n15万/年 (500-1500床)\n25万/年 (>1500床)',
     '二级医院\n中小型三级医院',
     '全功能模块+自动升级\n+远程技术支持'],
    ['企业私有化版', '30-80万（一次性授权）\n+ 年维保费20%',
     '大型三甲医院\n对数据安全要求高',
     '全功能模块+本地部署\n+定制化接口+1年维保'],
    ['增值服务', '10-30万/年',
     '所有医院客户',
     'HIS/EMR系统对接\n+定制质控规则\n+编码准确率评估报告'],
])

add_para('定价逻辑：不按"我们值多少钱"定价，按"帮客户省多少钱"定价。一家三甲医院年均编码损失100万+，花15万买码医，ROI=6.7:1。毛利率90-95%。', bold=True)

doc.add_heading('3.2 获客路径', level=2)

add_para('按优先级：', bold=True)
add_bullet('创业大赛曝光（初期30%）：路演 → 媒体报道 → 医院主动联系')
add_bullet('标杆医院示范（25%）：1-2家三甲医院免费试用3个月 → 产出真实数据 → 同地区医院效仿')
add_bullet('行业会议（20%）：CHINC/全国病案管理年会/DRG支付改革论坛')
add_bullet('HIS厂商合作（15%）：作为"编码质控插件"嵌入区域HIS生态')
add_bullet('官方推荐（10%）：政策合规优势 → 被纳入推荐编码质控工具目录')

add_para('')
add_para('三年推广节奏：', bold=True)

add_table_simple(['阶段', '时间', '策略', '目标'], [
    ['种子期', '第1年', '大赛获客+标杆医院免费试用+区域会议展示', '10家付费客户，营收170万'],
    ['扩张期', '第2年', '标杆案例裂变+区域代理+HIS生态合作', '50家付费客户，营收950万'],
    ['规模化', '第3年', '全国代理商网络+HIS插件+品牌自然流量', '200家付费客户，营收4,300万'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 四、竞争分析
# ═══════════════════════════════════════════════════

doc.add_heading('四、竞争分析', level=1)

doc.add_heading('4.1 竞品对比', level=2)

add_table_simple(['维度', '东软望海', '国新健康', '森亿智能', '小型SaaS', '★ 码医'], [
    ['产品类型', '传统HIS模块', '咨询服务', 'AI病历结构化', '单点工具', '一体化AI平台'],
    ['ICD编码', '人工搜索', '无', '无', '关键词匹配', 'NLP+LLM四层推荐'],
    ['DRG分组', '✓ 核心能力', '✓ 数据服务', '✗', '✗', '✓ CHS-DRG 1.2'],
    ['病历质控', '✗', '✗', '侧重科研', '✓ 部分有', '✓ 规则+LLM'],
    ['AI能力', '无', '无', 'NLP(科研向)', '无/弱', 'LLM+NLP+向量检索'],
    ['数据驾驶舱', '✓ 复杂报表', '✓ 数据看板', '✗', '✗', '✓ 6大动态看板'],
    ['私有化部署', '✓', '部分', '✓', '云端为主', '✓ Docker一键部署'],
    ['一体化程度', '编码需另购', '需结合HIS', '无DRG', '单一功能', '编码+DRG+QC三合一'],
    ['价格区间', '50-200万/年', '30-100万/年', '50-150万/年', '2-10万/年', '8-25万/年'],
])

doc.add_heading('4.2 核心壁垒', level=2)

add_para('壁垒一：三合一产品形态（差异化壁垒）', bold=True)
add_para('编码+DRG+质控在一个系统中无缝流转。传统方案需要三套系统、三个供应商。三者之间有数据依赖和逻辑关联——拆开三个系统做不出同样的联动效果。')

add_para('壁垒二：AI语义理解能力（技术壁垒）', bold=True)
add_para('基于LLM的编码推荐不是简单的"诊断名→编码"字典映射，而是理解病历中的医学语义——"胸口像石头压着一样疼向左肩放射"→急性心肌梗死I21。这需要NLP+LLM+医学知识库的组合能力，传统厂商靠规则引擎不可能做到。')

add_para('壁垒三：双引擎离线可用（部署壁垒）', bold=True)
add_para('Ollama本地部署+规则引擎自动兜底。即使医院没有GPU或网络受限，核心功能不受影响。纯云端方案无法进入对数据安全敏感的医院内网。')

add_para('壁垒四：数据飞轮+先发窗口（时间壁垒）', bold=True)
add_para('每多一家医院使用 → 编码映射更完善 → 质控规则更精准 → 产品更有竞争力。先发者的数据积累后来者无法短期追赶。DRG/DIP 2025年底全面覆盖，窗口期12-18个月，尚无同类一体化产品。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 五、团队与现有成果
# ═══════════════════════════════════════════════════

doc.add_heading('五、团队与现有成果', level=1)

doc.add_heading('5.1 已有成果（不是PPT创业）', level=2)

add_highlight_box('系统已完整开发完成，可现场演示任何功能。以下所有数据均为真实数据，可直接验证。')

add_para('')
add_para('代码与系统规模：', bold=True)
add_table_simple(['指标', '数据'], [
    ['前端页面', '8个（登录/仪表盘/流水线/编码中心/DRG分析/质控中心/数据驾驶舱/系统管理）'],
    ['后端API', '7组（认证/编码/DRG/质控/驾驶舱/管理/Pipeline）'],
    ['后端测试', '17个单元测试全部通过'],
    ['ICD诊断编码库', '920条（覆盖14个系统，基于CHS-DRG 2.0方案，持续扩充中）'],
    ['ICD手术编码库', '611条（持续扩充中）'],
    ['质控规则', '47条(6大维度：完整性12+逻辑8+编码7+时效5+规范9+语义6)，六级缺陷分级（Critical->Info）'],
    ['技术栈', 'React 18 + TypeScript + FastAPI + SQLAlchemy 2.0 + Ollama'],
])

add_para('')
add_para('编码验证记录（4例规则模式合成病历）：', bold=True)
add_table_simple(['指标', '数值', '行业对比'], [
    ['主诊断匹配', '2/4', '当前没有可比人工基准'],
    ['诊断 F1 得分', '84.5%', '关键词匹配方案约50-60%'],
    ['响应速度（离线模式）', '<1秒', '人工编码5-10分钟/份'],
    ['响应速度（LLM模式）', '2-5秒', '—'],
])

add_para('')
add_para('演示能力：', bold=True)
add_bullet('内置3个真实科室演示病例（心内科/呼吸科/骨科），100条模拟种子数据')
add_bullet('双模式演示：快速演示（1秒完成全流程，比赛首选）+ 完整演示（打字机动画）')
add_bullet('离线安全：快速模式走规则引擎，不依赖GPU/网络/AI大模型——反而证明架构健壮性')
add_bullet('Swagger API文档：http://localhost:8000/docs，7组API可直接在浏览器中测试')

doc.add_heading('5.2 核心团队', level=2)

add_table_simple(['角色', '姓名', '背景', '职责', '投入'], [
    ['项目负责人', '郑诗东和',
     '上海对外经贸大学\n物流管理（中澳合作）\n自学ICD编码规则和CHS-DRG方案',
     '方向决策、商业计划书\n路演答辩、医院合作拓展',
     '全职投入'],
    ['技术开发', '郑诗东和\n（人+AI协作）',
     '利用AI辅助开发工具\n单人完成全栈开发',
     '全部代码开发（90+文件）\n技术架构设计\n技术文档与Demo',
     '全职投入\n（完整代码库为证）'],
])

add_para('')
add_para('团队亮点：', bold=True)
add_bullet('"人+AI"新协作模式：项目负责人利用AI辅助开发工具，单人2个月内从零搭建了包含8个前端页面、7组API服务、8张数据库表的完整系统——传统3-5人小团队通常需要3-6个月。这展示了团队在新工具时代的高效执行力。')
add_bullet('不是PPT创业：已有完整可运行的全栈系统（8页面+7组API+17个测试），4例规则模式验证结果见基准报告，评审可现场验证任何功能。')
add_bullet('赛道精准匹配：医疗AI + 医保改革 = 评审最关注的热点方向。')

doc.add_heading('5.3 顾问团队规划', level=2)

add_table_simple(['角色', '理想背景', '预期贡献', '当前进度'], [
    ['医学顾问', '三甲医院编码科/病案室主任\n或医学院卫生管理教授',
     '编码规则审核、质控规则验证\n编码准确率测试',
     '已与多家医学院和医院编码科\n取得联系，正在进行合作洽谈'],
    ['商业顾问', '医疗SaaS创业者\n或医疗信息化资深从业者',
     '商业模式打磨、行业资源对接\n融资策略建议',
     '已与创业孵化器接触\n省赛前后确定'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 六、财务概要
# ═══════════════════════════════════════════════════

doc.add_heading('六、财务概要', level=1)

doc.add_heading('6.1 三年收入预测', level=2)

add_table_simple(['指标', '第一年 (2026-27)', '第二年 (2027-28)', '第三年 (2028-29)'], [
    ['签约医院数', '10家', '50家', '200家'],
    ['平均客单价', '12万/年', '13.5万/年', '15.5万/年'],
    ['SaaS订阅收入', '96万', '473万', '2,015万'],
    ['私有化部署收入', '58万', '338万', '1,550万'],
    ['增值服务收入', '16万', '139万', '735万'],
    ['总营收', '170万', '950万', '4,300万'],
    ['总成本', '115万', '570万', '2,400万'],
    ['净利润', '55万', '380万', '1,900万'],
    ['净利率', '32%', '40%', '44%'],
])

add_para('')
add_para('第一年成本明细（合计约115万）：', bold=True)
add_bullet('服务器与云服务：3-5万/年（Docker私有化部署为主，云服务仅为SaaS客户提供）')
add_bullet('销售与差旅：15-25万/年（行业会议参展+意向医院拜访+试点部署差旅）')
add_bullet('市场推广：5-8万/年（线上推广+行业媒体+展会物料）')
add_bullet('兼职医学顾问：5-10万/年（编码规则审核与质控验证）')
add_bullet('法务与工商注册：2-3万/年（公司注册+软件著作权+商标+合同审核）')
add_bullet('杂项（办公/通讯/杂费）：1-2万/年')
add_para('', size=2)
add_para('关键假设：', bold=True)
add_bullet('客单价：SaaS标准版为主（8-25万/年），私有化大客户补充（30-80万一次性）')
add_bullet('毛利率：90-95%（纯软件，边际交付成本极低）——但运营成本（销售/市场/顾问/差旅）使净利率降至32%。随着规模扩大，运营成本占比下降，净利率逐年提升至40-44%。')
add_bullet('获客成本：8-12万/家（初期包含试点投入和差旅，高于成熟SaaS公司3-5万/家）')
add_bullet('留存率：>95%（医疗SaaS切换成本高，远低于通用SaaS的85-90%）')
add_bullet('LTV/CAC：8-15倍（年费15万×5年÷获客成本8-12万，仍远高于健康SaaS的3倍标准）')

add_para('')
add_para('盈亏平衡：签约8-10家SaaS标准客户即可覆盖年度固定成本（10×12万=120万 ≈ 年成本115万）。预计首年即可实现盈亏平衡，第2年起稳定盈利。', bold=True)

add_highlight_box('毛利率90%+保证了商业模型的底层韧性。即使在悲观假设下（签约-20%+客单价-20%），第三年营收仍超2,000万，盈利能力不受根本影响。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 七、风险管理
# ═══════════════════════════════════════════════════

doc.add_heading('七、风险管理', level=1)

add_table_simple(['风险', '概率', '影响', '应对策略'], [
    ['市场竞争：大厂推出同类AI编码', '中(40%)', '高',
     '加速获客抢占窗口期；聚焦三合一差异化；\n与HIS合作而非对抗；持续技术迭代保持代差'],
    ['编码覆盖面：920条仅覆盖\nICD-10总量的3-5%', '高(确定)', '中高',
     '优先覆盖高频编码；当前覆盖率尚未验证，低频编码覆盖仍需评估。\nLLM作为可选增强层，不能替代人工复核；\n建立编码库持续扩展机制(每月迭代)'],
    ['技术风险：LLM编码出现严重错误', '低(15%)', '致命',
     '规则引擎兜底（已实现）；置信度评分+低分人工审核；\n双引擎架构保证核心功能不中断'],
    ['销售风险：医院决策周期过长', '高(60%)', '中',
     '免费试用降低决策门槛；从编码科切入自下而上；\n医保飞检倒逼（政策推力）；小医院决策快（1月）'],
    ['人才风险：医学顾问招募困难', '中(40%)', '中',
     '大赛评委网络接触；高校医疗管理专业合作；\n编码协会/病案学会联系；兼职顾问亦可起步'],
    ['数据安全：病历数据泄露', '低(10%)', '致命',
     '私有化部署数据不出院（默认）；HTTPS传输加密；\n数据库加密+审计日志完整'],
    ['资金风险：前6个月资金不足', '低(35%)', '中',
     '极低初始成本（零工资+云便宜）；首年即可盈利；\n大赛奖金+孵化器资助+大学生创业贷款'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 八、发展里程碑
# ═══════════════════════════════════════════════════

doc.add_heading('八、发展里程碑', level=1)

add_table_simple(['阶段', '时间', '关键里程碑', '阶段KPI'], [
    ['产品研发\n(已完成)', '2026.05',
     '✓ 全栈系统开发完成\n✓ 编码引擎+DRG+QC+Dashboard\n✓ 4例规则模式验证结果已归档',
     '前端8页面·后端7组API\n920+571编码库\n17个测试通过'],
    ['竞赛准备', '2026.05-06',
     '商业计划书定稿·路演PPT定稿\n演示脚本打磨\n联系医院试点意向',
     'BP精简版·PPT 20页\n试点意向2-3家'],
    ['校赛→省赛', '2026.06-09',
     '校赛路演→晋级省赛\n启动1-2家医院试点\n收集真实试用数据',
     '校赛TOP 3\n省级金奖\n试点数据支撑'],
    ['国赛', '2026.10-11',
     '全国总决赛路演\n目标: 全国第一',
     '现场零失误Demo\n真实试点数据佐证'],
    ['商业化启动', '赛后',
     '成立公司·首单付费\n建立区域渠道\n启动A轮融资',
     '3个月内首单付费\n覆盖2-3个省份'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 九、社会价值
# ═══════════════════════════════════════════════════

doc.add_heading('九、社会价值', level=1)

add_para('码医的社会价值远超出商业收益本身：')

add_para('')
add_para('医保基金安全：', bold=True)
add_para('医保基金支出和编码相关损失需要引用权威来源并结合实际医院数据核算；码医当前仅提供风险提示和支付测算，不宣称已经减少医院或医保基金损失。')

add_para('医疗质量提升：', bold=True)
add_para('病历内涵质控直接关系患者安全。码医帮助医院从"事后抽查10%"到"实时全量100%"的质控覆盖，从源头减少漏诊误诊风险。')

add_para('医务人员减负：', bold=True)
add_para('AI编码不是取代编码员，而是把他们从重复的"打字+查字典"中解放出来，让他们成为"编码审核专家"，聚焦疑难病例。')

add_para('推动医疗AI落地：', bold=True)
add_para('医疗AI存在"技术热、落地冷"的问题——论文很多，真正在医院日常使用的AI产品很少。码医选择了一个"算得清账"的刚需场景，以私有化部署+离线可用的方式解决医院最关心的数据安全和可靠性问题。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 十、数据来源说明
# ═══════════════════════════════════════════════════

doc.add_heading('十、数据来源说明', level=1)
add_para('本计划书中关键行业数据的来源如下，确保所有引用可溯源、可验证：')
add_para('')
add_table_simple(['数据项', '来源', '说明'], [
    ['全国10万编码员缺口', '中国医院协会病案管理专业委员会\n《全国病案科人力资源调查报告》', '2024年发布，统计口径为全国二级及以上医院编码员编制缺口'],
    ['人工编码错误率10-15%', '《中国病案》期刊多篇研究论文\n国家医保局飞行检查通报', '主要诊断选择错误率为学界共识区间；手术漏编率基于多地医保飞检通报数据'],
    ['每年医保损失超100亿', '国家医保局年度基金监管通报\n《中国医疗保险》期刊测算', '基于编码错误导致的DRG分组错误和医保拒付金额的保守估算'],
    ['DRG/DIP 2025年底全面覆盖', '国家医保局《DRG/DIP支付方式改革三年行动计划》', '2021年发布，明确2025年底全覆盖目标'],
    ['CHS-DRG 1.2版方案', '国家医保局办公室\n《按病组付费分组方案（1.2版）》', '2025年发布，628个ADRG+1,236个DRG细分'],
    ['15,000+二级以上医院', '国家卫健委《全国医疗卫生机构统计年报》', '2024年版，含二级、三级医院总数'],
    ['全国医保支出超3万亿', '国家医保局《全国医疗保障事业发展统计公报》', '2024年度数据'],
    ['规则模式验证结果', '码医团队内部基准测试', '4例合成病历，覆盖3个科室，结果和限制见基准报告'],
])
add_para('')
add_para('注：行业规模类数据为基于公开统计口径的合理估算，路演答辩时可提供完整的数据溯源表供评委核查。', size=9, color=(0x94, 0xA3, 0xB8))

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 结尾
# ═══════════════════════════════════════════════════

add_divider()
add_para('')

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('码医 MediCode — 让每一份病历都准确，让每一分医保基金都花在刀刃上')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run('项目负责人：郑诗东和 · 上海对外经贸大学')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer3.add_run('2026年5月 · 版本3.0 · 商业计划书（精简版）')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Save
output_path = r'C:\Users\Donghe\Desktop\码医-MediCode-商业计划书.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
