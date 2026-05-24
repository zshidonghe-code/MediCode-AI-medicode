"""Generate 码医-MediCode 演示脚本 .docx file."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(22)
    elif level == 2:
        h.font.size = Pt(16)
    elif level == 3:
        h.font.size = Pt(13)


def add_para(text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return p


def add_table_simple(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()


# ═══════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════

doc.add_heading('码医 MediCode — 路演演示脚本', level=1)

add_para('总时长：5分钟 | 版本：v1.0 | 日期：2026年5月24日', italic=True, size=10)

doc.add_heading('一、演示前准备清单', level=2)

add_para('1. 确保后端和前端都已启动（python -m src.main + npm run dev）', bold=False)
add_para('2. 浏览器打开 http://localhost:5173，登录 admin / admin123', bold=False)
add_para('3. 打开两个浏览器Tab：', bold=False)
add_para('    Tab 1: 智能流水线页面（/pipeline）— 核心演示页', bold=False)
add_para('    Tab 2: 数据驾驶舱页面（/dashboard）— 展示运营数据', bold=False)
add_para('4. 确认Dashboard有数据（系统启动时会自动种子100条演示数据）', bold=False)
add_para('5. 如果Ollama已安装并运行，侧边栏会显示"AI在线"（加分项）', bold=False)
add_para('6. 关闭所有通知和不必要的窗口，确保屏幕干净', bold=False)

doc.add_heading('二、5分钟演示脚本', level=2)

doc.add_heading('段落1：问题冲击（0:00-0:30）', level=3)

add_table_simple(['时间', '画面', '话术'], [
    ['0:00', '打开码医系统，停在登录页，展示品牌Logo',
     '各位评委老师好，我是郑诗东和，来自上海对外经贸大学。今天我要介绍的项目是——码医，一个用AI帮助医院解决DRG编码错误的系统。'],
    ['0:10', '切换到Dashboard页面，指向统计数字',
     '先给大家看一组数据——全国10万编码员缺口，人工编码错误率10-15%。2025年底DRG/DIP改革全面覆盖，编码直接决定医院收入。'],
    ['0:20', '指向"预估收入"和"优化空间"对比',
     '一家三甲医院，每年因为编码错误损失的医保收入超过100万。全国3,200家三级医院，加起来就是数十亿的浪费。这不是管理问题，是工具问题。'],
])

doc.add_heading('段落2：产品演示 — 智能流水线（0:30-2:00）', level=3)

add_table_simple(['时间', '画面', '话术'], [
    ['0:30', '切换到Pipeline页面，点击"演示模式"',
     '接下来我给大家演示码医的核心能力。我打开"演示模式"，系统会自动模拟一份真实病历的输入过程。'],
    ['0:35', 'Typewriter打字机效果开始，病历逐字出现',
     '大家可以看到，AI正在自动输入一份病历——这是一位65岁男性，急性心肌梗死患者的出院小结。'],
    ['0:55', '打字完成，系统自动开始第1步"智能编码"',
     '打字完成后，系统自动进入NLP智能编码。从病历文本中提取诊断实体，匹配ICD-10编码。'],
    ['1:05', '编码结果展示：主诊断I21.900 + 次要诊断 + 手术操作',
     '大家看——主要诊断：急性心肌梗死I21.900，次要诊断：冠心病I25.100、高血压I10、糖尿病E11，手术操作：冠状动脉支架植入36.07。AI在1秒内完成了编码员需要10分钟才能完成的工作。'],
    ['1:10', '系统自动进入第2步"质控检查"',
     '接下来自动进入质控检查——系统用100多条规则扫描病历，检查有无遗漏、矛盾和不规范之处。'],
    ['1:20', '质控结果展示：评分 + 缺陷列表',
     '质控评分90分，发现1个一般缺陷——"缺少出院医嘱部分"。系统给出了具体建议。编码员可以点击采纳或忽略。'],
    ['1:25', '系统自动进入第3步"DRG分组"',
     '然后进入DRG分组。根据编码结果和患者信息，系统自动判定为FC1组——PCI+STEMI手术组，权重3.74。'],
    ['1:30', 'DRG结果展示 + 费用测算',
     '最后一步——费用测算。RW 3.74 × 费率12,000元 = 预估医保支付44,880元。'],
    ['1:35', '指向费用数字',
     '现在请大家注意——如果我漏编了次要诊断，DRG权重会降低，医院可能少收5,000-8,000元。一份病历差8,000，一年15万份病历就是120万。码医的价值就是帮医院守住这120万。'],
    ['1:50', '切换到Dashboard',
     '现在大家看到的是数据驾驶舱。刚才我们处理的这份病历，结果已经自动保存到数据库中，Dashboard实时更新。'],
])

doc.add_heading('段落3：数据驾驶舱（2:00-3:00）', level=3)

add_table_simple(['时间', '画面', '话术'], [
    ['2:00', 'Dashboard概览：总病例数、CMI、质控通过率',
     '这是模拟一家医院3个月的运营数据——100位患者，100份病历，全流程编码+QC+DRG。'],
    ['2:15', '向下滚动，展示科室排名',
     '科室编码排名——心内科CMI最高，骨科DRG权重最高。这是医院管理者最关心的数据。'],
    ['2:25', '展示质控趋势图',
     '质控合格率趋势——可以看到随着AI编码的推广，质控评分稳步上升。'],
    ['2:40', '展示收入分析图',
     '医保收入分析——实际收入 vs 优化预估。可以看到，通过AI编码优化，每家医院每年有15-30%的收入提升空间。'],
    ['2:55', '切回Pipeline页面',
     '这些数据都是我们系统真实生成的，不是PPT截图。'],
])

doc.add_heading('段落4：技术壁垒（3:00-3:45）', level=3)

add_table_simple(['时间', '画面', '话术'], [
    ['3:00', '切换到Swagger API文档（/docs）',
     '大家可能会问——市面上不是有很多医疗软件公司吗？我给大家看几个关键差异。'],
    ['3:05', '回到系统，指向侧边栏AI状态指示器',
     '第一，我们用的是AI大模型，不是传统规则引擎。可以看到侧边栏实时显示AI在线状态——我们用Ollama本地部署Qwen2.5大模型，数据不出医院内网。'],
    ['3:15', '展示编码结果的置信度评分',
     '第二，三层编码推荐策略——NLP实体识别、TF-IDF语义检索、LLM推理。不是简单的关键词匹配，而是真正理解医学语义。"胸口像石头压着一样疼"→它知道是心绞痛。'],
    ['3:25', '展示Pipeline全流程',
     '第三，也是最重要的差异化——三合一。编码+DRG分组+质控，一份病历进去，三个结果同时出来。市场上没有其他产品做到这一点。传统方案要买三套系统，找三个供应商，数据互不相通。'],
    ['3:40', '',
     '第四，私有化部署。整套系统可以在医院内网运行，Docker Compose一键部署。即使没有GPU、没有网络，核心功能也不受影响。'],
])

doc.add_heading('段落5：商业模式 + 愿景（3:45-4:30）', level=3)

add_table_simple(['时间', '画面', '话术'], [
    ['3:45', '回到PPT最后一页或口述',
     '商业模式方面——SaaS订阅，按医院床位数分级定价，8到25万一年。一家医院一年省100万，花15万买我们的系统，ROI超过6:1。这个账医院算得过来。'],
    ['4:00', '',
     '市场空间——全国3,200多家三级医院，10,000多家二级医院。保守估计可及市场22亿/年，长期45亿/年。'],
    ['4:10', '',
     '我们的目标路径：第一年签约10家医院，营收170万；第二年50家，营收950万；第三年200家，营收4,300万。'],
    ['4:20', '',
     '目前系统已完整开发完成并可以现场演示。校赛通过后，我们会立即启动1-2家医院的试点合作。'],
])

doc.add_heading('段落6：结尾（4:30-5:00）', level=3)

add_table_simple(['时间', '画面', '话术'], [
    ['4:30', '关闭浏览器，回到PPT结尾页',
     '最后回到我们项目的初心——'],
    ['4:35', 'PPT结尾页：让每一份病历都准确',
     'DRG/DIP付费改革是国家医疗改革的核心政策。编码的准确性，直接关系到医保基金的安全，关系到每一个病人的就医公平。'],
    ['4:45', '',
     '码医的目标是——让AI帮医院做对编码，让医保基金花在刀刃上，让医务人员从重复劳动中解放出来。'],
    ['4:55', '鞠躬',
     '我是郑诗东和，码医MediCode。谢谢各位评委老师！'],
])

doc.add_heading('三、评委可能的提问 & 回答准备', level=2)

add_para('', size=8)

add_table_simple(['可能的问题', '回答要点'], [
    ['"你们的技术壁垒在哪里？"',
     '三层AI架构 + 三合一产品形态 + 私有化部署 + 数据飞轮。市场上没有同类一体化产品，窗口期12-18个月。'],
    ['"和东软望海/国新健康比有什么优势？"',
     '他们是传统HIS/咨询服务商，用的是老规则引擎，没有AI语义理解能力，产品割裂。我们是AI驱动的一体化平台。'],
    ['"编码准确率到底多少？"',
     '目前系统通过300+诊断编码库和4层推荐策略，计划通过医院试点验证目标95%+。我们有LLM增强和规则兜底机制。'],
    ['"怎么赚钱？什么时候盈利？"',
     'SaaS年费+私有化授权，毛利率90%，首年即盈利。客户ROI超过6:1。详细财务预测见商业计划书。'],
    ['"团队能做出来吗？"',
     '已经有完整可运行的系统，不是PPT创业。后端7组API，前端8个页面，17个单元测试通过，现场可以演示任何功能。'],
    ['"为什么是你来做？别人为什么不做？"',
     'AI+医疗编码是交叉领域——需要同时懂NLP/LLM技术又懂ICD/DRG医学知识。传统HIS厂商缺AI能力，AI创业公司缺医疗领域知识。我们的AI Agent模式填补了这个缺口。'],
    ['"有医院合作吗？"',
     '目前处于产品完成、寻找试点的阶段。校赛后计划联系1-2家医院开展免费试用，用真实数据迭代优化。'],
    ['"数据安全怎么保证？"',
     '全私有化部署，病历数据不出医院内网。支持离线运行。不使用外部API（如ChatGPT），Ollama本地推理。'],
])

doc.add_heading('四、备用方案：如果出问题怎么办', level=2)

add_table_simple(['故障场景', '应急方案'], [
    ['Ollama不可用（AI离线）', '系统自动回退到规则引擎，功能不受影响。侧边栏显示"AI离线"，'
     '如实告诉评委这是设计好的兜底机制——反而证明我们的双引擎架构是真实的。'],
    ['前端或后端崩溃', '重新运行启动命令。提前准备好备用终端窗口，后端和前端都已在后台运行。'],
    ['浏览器卡顿', '提前清理浏览器缓存，关闭多余Tab，确保Chrome/Edge最新版。'],
    ['网络断开', '系统完全离线可用（SQLite + 规则引擎），断网不影响Demo。'],
    ['投影仪不兼容/分辨率问题', '自带笔记本+HDMI转换器，提前测试投影。系统使用16:9比例，兼容大部分投影仪。'],
    ['演示病例数据丢失', '系统启动自动种子数据，重启后端即可恢复。Dashboard始终有100条演示数据保底。'],
])

doc.add_heading('五、演示模式速查', level=2)

add_para('Pipeline演示模式操作：', bold=True)
add_para('1. 页面顶部点击"演示模式"按钮')
add_para('2. 下拉菜单可选择病例：心内科·急性心肌梗死 / 呼吸科·COPD急性加重 / 骨科·股骨颈骨折')
add_para('3. 速度调节：快（12ms/字）/ 中（28ms/字）/ 慢（55ms/字），一般选中速')
add_para('4. 演示自动运行：打字完成 → 自动分析 → Confetti庆祝 → 可点"重新演示"重播')
add_para('5. 可在自动演示期间点击"停止"中断，手动操作')

add_para('')
add_para('Dashboard展示路径：', bold=True)
add_para('1. 左侧菜单→数据驾驶舱')
add_para('2. 顶部概览区展示4个核心KPI卡片')
add_para('3. 滚动查看：科室排名 → 质控趋势 → 编码准确率 → 高频缺陷 → 收入分析')
add_para('4. 日期范围选择器可切换不同时间段')

add_para('')
add_para('Swagger API文档展示：', bold=True)
add_para('1. 侧边栏底部"API 文档"链接，或直接打开 http://localhost:8000/docs')
add_para('2. 展示7组API路由（auth/coding/drg/qc/dashboard/admin/pipeline）')
add_para('3. 可选：在Swagger页面直接尝试调用一个API端点', italic=True)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('码医 MediCode — 让每一份病历都准确，让每一分医保基金都花在刀刃上')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Save
output_path = r'C:\Users\Donghe\Desktop\码医-MediCode-演示脚本.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
