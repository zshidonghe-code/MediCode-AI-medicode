"""Generate 码医-MediCode project introduction .docx file."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# Base style
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(20)
        h.paragraph_format.space_after = Pt(8)
    elif level == 3:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, italic=False, size=None, color=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    return p


def add_code_block(code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)
        shading = p._element.get_or_add_pPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '1E293B'
        })
        shading.append(shd)


def add_table_simple(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))


# ═══════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════

# Title
title = doc.add_heading('码医 MediCode', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI医疗DRG编码与病历质控系统')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# ── 1. 项目概况 ──
doc.add_heading('一、项目概况', level=2)

add_para('码医（MediCode）是一个面向医院编码科和质控科的AI工作台。其核心能力是：输入一份住院病历，AI自动完成 ICD编码 → DRG分组 → 质控检查 → 费用测算，全流程一体化输出。')

add_para('项目的直接价值是帮医院避免医保拒付（编码错了就扣钱），深层价值是让医保基金的每一分钱都花在刀刃上。')

add_table_simple(['项目', '内容'], [
    ['定位', 'SaaS + 私有化部署的医院智能化工具'],
    ['团队', '郑诗东和（项目负责人）+ Claude（AI技术负责人）'],
    ['赛道', '互联网+大学生创新创业大赛（主攻国赛第一）'],
    ['状态', 'MVP核心功能已开发完成，进入产品打磨阶段'],
    ['代码路径', 'C:\\Users\\Donghe\\Desktop\\码医-MediCode\\'],
])

# ── 2. 问题背景 ──
doc.add_heading('二、问题背景', level=2)

doc.add_heading('2.1 DRG/DIP付费改革是国家级政策', level=3)
add_para('2021年国务院办公厅发文明确全面推行DRG/DIP付费。2025年底前覆盖所有符合条件的住院医疗机构。这意味着全国每一份住院病历的ICD编码，将直接决定医院的医保收入。不再是"看好病就能收钱"，而是"编对码才能收钱"。')

doc.add_heading('2.2 医院面临的四大痛点', level=3)

add_table_simple(['痛点', '严重程度', '具体表现'], [
    ['编码员严重不足', '★★★★★',
     '全国编码员缺口超10万人。一个三甲医院年均出院15万人次，需要10+名专职编码员，实际往往只有3-5人。每人每天处理50-80份病历，高强度下错误不可避免。'],
    ['编码错误率居高不下', '★★★★★',
     '人工编码错误率10-15%。编码错误直接导致DRG分组错误，进而医保拒付。一个主要诊断编码选错，医院直接损失几千到几万元。更严重的是——医保飞检倒查3年，一个错误可能被追索3年的费用差额。'],
    ['病历内涵质控缺失', '★★★★☆',
     '医生写病历普遍不规范：诊断和手术部位矛盾、出院小结缺项、症状当主诊断。传统质控靠人工抽查，覆盖率不到10%，大量问题病历被遗漏，成为飞检时的定时炸弹。'],
    ['医保飞检罚款', '★★★★★',
     '医保飞检常态化，编码问题是重点检查项。一次飞检查出编码问题，罚款动辄数十万到数百万。医院不是不愿改正，而是缺少有效的自查工具。'],
])

doc.add_heading('2.3 为什么AI能解决这个问题', level=3)

add_bullet('ICD编码本质是"文本→编码"的映射问题。病历描述"胸口压榨性疼痛向左肩放射"应编码为I20.0（不稳定型心绞痛）。这恰恰是NLP+LLM最擅长的事。')
add_bullet('DRG分组是有明确规则的逻辑推理。CHS-DRG 1.2版公布了完整的MDC→ADRG→DRG分组方案，可以通过规则引擎100%代码化。')
add_bullet('80%以上的质控问题可以用规则+语义理解覆盖。比如"手术记录写了胆囊切除术但诊断没提胆囊"——用NLP比对手术名和诊断名就能发现。')

# ── 3. 产品功能 ──
doc.add_heading('三、产品功能', level=2)

doc.add_heading('3.1 智能ICD编码引擎', level=3)
add_para('输入：住院病历全文（入院记录、病程记录、手术记录、出院小结等）', bold=True)
add_para('输出：主要诊断编码（1个）+ 其他诊断编码（多个）+ 手术操作编码（多个），每个编码附置信度评分。', bold=True)

add_para('工作原理 —— 4层推荐策略：', bold=True)
add_bullet('第1层：NLP实体识别 → 从病历文本中提取诊断名、手术名、部位、药物等医学实体')
add_bullet('第2层：数据库精确匹配 → 在已有ICD编码库中精确查找')
add_bullet('第3层：语义向量检索 → TF-IDF + char n-gram (1-3) + 余弦相似度，解决"胸口疼"→"心绞痛"这类同义不同表述的匹配问题')
add_bullet('第4层：LLM推理 → Ollama (qwen2.5:7b) 对候选编码做最终推荐，处理复杂疑难多诊断关联场景')

add_para('技术指标：', bold=True)
add_bullet('诊断映射库：180+条目（覆盖循环/呼吸/消化/肝胆/泌尿/内分泌/神经/骨骼/肿瘤/感染/血液/皮肤/五官/产科共14个系统）')
add_bullet('手术映射库：60+条目（覆盖心血管/骨科/普外/胸科/泌尿/妇产/神外/五官/血管共9个专科）')
add_bullet('语义检索：TF-IDF字符级1-3 gram，5000特征维度')
add_bullet('目标准确率：主要诊断≥95%，全部诊断≥90%，手术操作≥92%')

doc.add_heading('3.2 DRG自动分组器', level=3)
add_para('输入：ICD编码结果 + 患者基本信息（年龄、性别、住院天数、离院方式）', bold=True)
add_para('输出：MDC大类 → ADRG核心分组 → DRG细分编码 → 权重(RW) + 费率 + 预估支付金额 + CC/MCC标记', bold=True)

add_para('核心流程：', bold=True)
add_bullet('ICD编码 → MDC大类判定（根据主要诊断映射到26个MDC之一）')
add_bullet('手术/非手术分组判定（是否含手术操作编码）')
add_bullet('ADRG细分组匹配（CHS-DRG 1.2版628个核心分组）')
add_bullet('CC/MCC逻辑判定（次要诊断是否有合并症或严重合并症，影响支付权重）')
add_bullet('最终DRG确定 + 权重查表 + 支付金额计算 = RW × 费率')

doc.add_heading('3.3 病历内涵质控引擎', level=3)
add_para('输入：病历全文 + 编码结果', bold=True)
add_para('输出：质控缺陷清单，每条包含规则编号、严重程度、问题描述、改进建议。', bold=True)

add_table_simple(['规则类型', '示例规则', '检查方式'], [
    ['完整性检查', 'QC-101：出院小结必需段落完整性', '规则引擎'],
    ['逻辑一致性', 'QC-102：手术与诊断一致性（如"阑尾切除术"对应有"阑尾炎"诊断）', 'LLM驱动'],
    ['逻辑一致性', 'QC-103：主要诊断选择正确性（不能把症状当主诊断）', 'LLM驱动'],
    ['编码一致性', 'QC-201：编码与诊断文本匹配检查（诊断写"左侧"但编码没体现部位）', 'LLM驱动'],
    ['编码一致性', 'QC-202：漏编次要诊断检查（病历提了"高血压"但编码里没有）', 'LLM驱动'],
    ['时效性检查', 'QC-301：入院记录是否在入院后24小时内完成', '规则引擎'],
    ['时效性检查', 'QC-302：手术记录完成时间检查', '规则引擎'],
    ['规范表达', 'QC-401：主要诊断不能为"待查"等不确定表述', '规则引擎'],
    ['规范表达', 'QC-402：诊断名称非正式化检测（如"心梗"应为"急性心肌梗死"）', '规则引擎'],
])

add_para('严重程度分级：CRITICAL（严重/扣100分）→ MAJOR（重要/扣50分）→ MINOR（一般/扣20分）→ INFO（提示/扣5分）')

doc.add_heading('3.4 数据驾驶舱', level=3)
add_para('面向医院管理层的运营分析看板，提供6类数据指标：')
add_bullet('全院DRG运营概览（总病例数、AI编码率、质控通过率、CMI均值）')
add_bullet('科室排名（编码数量、准确率、CMI值对比，金银铜牌标记）')
add_bullet('质控合格率趋势（日/周/月维度，可切换7/30/90/180天）')
add_bullet('AI vs 人工编码准确率对比趋势')
add_bullet('高频质控缺陷类型分布（12类问题）')
add_bullet('医保收入分析（实际收入、优化预估、优化空间，含12个月趋势图）')

doc.add_heading('3.5 系统管理', level=3)
add_para('管理员专属的数据管理功能：')
add_bullet('数据预览：实时显示各表记录数（患者/病历/编码/质控），统计卡片展示')
add_bullet('数据重置：支持 dry-run 预览模式（confirm=false），需输入 RESET 二次确认后执行，按外键安全顺序删除')
add_bullet('数据导出：支持编码结果、患者摘要、质控结果三种导出类型，可选 JSON 或 CSV 格式')
add_bullet('SQLite自增计数器重置：删除数据后自动重置自增ID，保证下次种子数据ID一致性')
add_bullet('权限控制：仅管理员角色可见，前后端双重鉴权')

doc.add_heading('3.6 智能编码流水线（核心演示页）', level=3)
add_para('这是面向比赛路演的一站式可视化页面。一条流水线串联全部流程：')

add_code_block('''[粘贴病历] -> [NLP智能编码] -> [质控检查] -> [DRG分组] -> [费用测算]
    |                                                       |
支持手动输入 / 文件上传                        动画计数器显示预计收入''')

add_para('')
add_para('演示模式（专为路演设计）：', bold=True)
add_bullet('内置3个真实病历样本（心内科心肌梗死、呼吸科COPD急性加重、骨科股骨颈骨折）')
add_bullet('Typewriter打字机效果：AI逐字自动输入病历内容，模拟真人打字')
add_bullet('自动触发全流程分析：打字完成后自动依次走完编码→质控→DRG→费用4步')
add_bullet('速度可调：快/中/慢三档，适应不同演示节奏')
add_bullet('完成后confetti五彩纸屑庆祝动画')
add_bullet('支持一键重播，可快速切换病例')
add_bullet('流水线结果自动持久化到数据库，Dashboard实时反映')
add_bullet('LLM在线状态指示器：侧边栏实时显示Ollama连接状态（在线/离线）')
add_bullet('文件上传支持 .txt / .docx / .pdf 三种格式')
add_bullet('QC采纳/忽略操作持久化到后端数据库，刷新不丢失')

# ── 4. 技术架构 ──
doc.add_heading('四、技术架构', level=2)

doc.add_heading('4.1 整体架构图', level=3)

add_code_block('''+-------------------------------------------------------+
|                  前端 (React 18 + TypeScript)          |
|  流水线 | 编码台 | DRG分组 | 质控中心 | 驾驶舱 | 登录  |
|                  (Ant Design 5 + ECharts)              |
+--------------------------+----------------------------+
                           | HTTP REST API
+--------------------------+----------------------------+
|                 后端 API 网关 (FastAPI)                 |
|  /api/v1/coding  |  /api/v1/drg  |  /api/v1/qc       |
|  /api/v1/dashboard  |  /api/v1/auth                  |
+-----+--------------+--------------+-------------------+
      |              |              |
      v              v              v
+----------+  +-----------+  +------------+
| NLP引擎   |  | ICD编码器 |  | 质控引擎    |
| -SOAP拆分 |  | -4层推荐  |  | -15条规则  |
| -实体识别 |  | -向量检索 |  | -LLM检查   |
|           |  | -LLM推荐  |  | -评分计算  |
+-----+----+  +-----+-----+  +------+-----+
      |             |               |
      +-------------+---------------+
                    |
                    v
+-------------------------------------------------------+
|                      数据层                            |
|   SQLite(开发) | Ollama qwen2.5:7b | TF-IDF向量索引   |
|   RuleBased规则引擎（离线兜底）                        |
+-------------------------------------------------------+''')

doc.add_heading('4.2 技术选型', level=3)

add_table_simple(['层级', '技术', '说明'], [
    ['前端框架', 'React 18 + TypeScript', '类型安全，生态成熟'],
    ['构建工具', 'Vite', '比Webpack快10倍，HMR秒级'],
    ['UI组件库', 'Ant Design 5.x', '企业级组件库，适合医疗产品'],
    ['图表', 'ECharts', '数据可视化，支持动态交互'],
    ['状态管理', 'Zustand', '轻量级替代Redux'],
    ['HTTP客户端', 'Axios', '拦截器支持，JWT自动续期'],
    ['后端框架', 'FastAPI (Python 3.14)', '异步支持，自动API文档生成'],
    ['ORM', 'SQLAlchemy 2.0', '异步引擎，成熟稳定'],
    ['数据库', 'SQLite（开发）/ PostgreSQL（生产）', '开发阶段零配置，生产可切换'],
    ['LLM推理', 'Ollama + qwen2.5:7b', '本地部署，数据不出医院内网'],
    ['向量检索', 'sklearn TF-IDF', 'char级n-gram(1-3)，中文语义匹配'],
    ['LLM编排', '自研双后端引擎', 'Ollama优先，RuleBased自动兜底'],
])

doc.add_heading('4.3 LLM引擎架构（亮点设计）', level=3)

add_code_block('''LLMEngine (统一接口)
    |
    +-- _get_backend() -> 自动检测
    |
    +-- OllamaBackend (首选)
    |     +-- HTTP -> localhost:11434
    |     模型: qwen2.5:7b
    |     用途: 编码推荐、质控检查、DRG优化
    |
    +-- RuleBasedBackend (兜底)
          +-- 内置规则 + 关键词匹配
             离线可用，无GPU需求''')

add_para('')
add_para('这个双后端设计的意义在于：当Ollama可用时，享受LLM的语义理解能力；当Ollama不可用时（比如医院服务器没有GPU），系统依然能正常运行，所有功能不中断。')

doc.add_heading('4.4 数据库设计（核心8张表）', level=3)

add_table_simple(['表名', '用途', '核心字段'], [
    ['patients', '患者主表', 'patient_id, name, gender, age, admission_date, discharge_date'],
    ['medical_records', '病历记录', 'record_type, content, department, doctor, created_at'],
    ['icd_codes', 'ICD编码库', 'code, name, category, version, py_code（拼音码）'],
    ['coding_results', '编码结果', 'coder_type (AI/人工), codes_json, confidence'],
    ['drg_groups', 'DRG分组定义', 'code, name, mdc, adrg, weight, rate'],
    ['qc_rules', '质控规则', 'rule_name, rule_type, severity, logic_expression'],
    ['qc_results', '质控结果', 'record_id, rule_id, line_number, snippet, status'],
    ['coding_logs', '编码审计日志', 'record_id, version, changes_json, operator, timestamp'],
])

# ── 5. 开发进度 ──
doc.add_heading('五、当前开发进度', level=2)

doc.add_heading('5.1 已完成', level=3)

add_para('后端（backend/src/）：', bold=True)
add_bullet('FastAPI应用骨架，CORS配置，lifespan事件管理（含自动种子数据 + LLM预热）')
add_bullet('8张数据表模型（SQLAlchemy 2.0 async）+ 17个单元测试全部通过')
add_bullet('7组API路由：auth、coding、drg、qc、dashboard、admin（管理）、pipeline（保存）')
add_bullet('NLP引擎：SOAP段落拆分 + 正则医学实体识别（诊断名、手术名、部位、药物）')
add_bullet('ICD编码器：4层推荐策略，300+诊断映射（JSON数据源统一），113手术映射')
add_bullet('DRG分组器：26个MDC大类判定 + CC/MCC合并症逻辑 + 分组接口 + 费用计算')
add_bullet('质控引擎：规则引擎 + LLM语义双模，15条规则，异步执行')
add_bullet('LLM引擎：OllamaBackend + RuleBasedBackend双后端，自动检测切换 + 健康检查端点')
add_bullet('向量搜索引擎：TF-IDF + char n-gram(1-3) + 余弦相似度语义检索')
add_bullet('Dashboard API：6个端点全部动态化，基于真实数据库查询')
add_bullet('Admin端点：数据重置（dry-run + 确认）+ 三种数据导出（JSON/CSV）')
add_bullet('Pipeline保存端点：流水线结果持久化到数据库，QC采纳/忽略端点实现数据库更新')
add_bullet('文件解析：支持 .txt / .docx / .pdf 三种格式自动解析')

add_para('前端（frontend/src/）：', bold=True)
add_bullet('Vite + React 18 + TypeScript项目骨架，自定义主题')
add_bullet('8个页面：Login / Pipeline / Coding / DRG / QC / Dashboard / Guide / Admin')
add_bullet('NotFoundPage：404通配路由，友好提示页')
add_bullet('AppLayout布局（渐变侧边栏 + LLM状态指示器 + API文档链接 + 品牌Logo + 用户菜单）')
add_bullet('AdminRoute角色守卫：非管理员访问 /admin 自动重定向')
add_bullet('LoginPage：深色渐变背景 + 浮动医疗图标 + 特征标签 + 团队页脚')
add_bullet('PipelinePage（核心）：4步流水线 + 文件上传(.txt/.docx/.pdf) + AnimatedCounter + 演示模式')
add_bullet('DRGPage：患者年龄/性别输入控件，支持自由填写')
add_bullet('DashboardPage：6个ECharts图表 + 日期范围筛选 + 动态数据')
add_bullet('AdminPage：数据预览统计 + 重置确认 + JSON/CSV导出下载 + Blob错误处理')
add_bullet('GuidePage：新手使用指南，大白话介绍系统使用方法')
add_bullet('Zustand认证管理 + Axios拦截器 + JWT流程 + pipelineAPI自动保存')

add_para('基础设施：', bold=True)
add_bullet('Docker Compose配置（PostgreSQL + Redis + Backend + Frontend）')
add_bullet('前后端Dockerfile')
add_bullet('Ollama已安装，qwen2.5:7b模型已下载')

doc.add_heading('5.2 待完成（按优先级）', level=3)

add_table_simple(['优先级', '任务', '预计耗时'], [
    ['P0', 'DRG完整分组匹配（628 ADRG -> 1236 DRG）', '2天'],
    ['P1', '质控规则扩充（从15条到50+条）', '3天'],
    ['P1', '编码准确率系统测试（100+份真实病历）', '2天'],
    ['P2', 'PostgreSQL迁移（替代SQLite）', '1天'],
    ['P2', '大屏演示分辨率适配', '0.5天'],
    ['P2', '电商化SaaS订阅后台（多租户、计费）', '5天'],
    ['P3', '医院HIS/电子病历系统对接接口', '待定'],
    ['P3', '医院试点合作', '待定'],
])

# ── 6. 商业模式 ──
doc.add_heading('六、商业模式', level=2)

doc.add_heading('6.1 收入模型', level=3)

add_table_simple(['模式', '定价逻辑', '特点'], [
    ['SaaS订阅', '500床以下 8万/年 | 500-1500床 15万/年 | 1500床以上 25万/年', '续费制，现金流稳定'],
    ['私有化部署', '一次性授权费 30-80万 + 年维保费20%', '数据不出院，符合医院安全要求'],
])

doc.add_heading('6.2 市场空间', level=3)
add_bullet('全国三级医院 3,200+ 家，二级医院 10,000+ 家')
add_bullet('理论市场规模：3,000家 × 平均15万/年 ≈ 45亿/年')
add_bullet('初期目标：第一年10家 → 第二年50家 → 第三年200家')

doc.add_heading('6.3 医院为什么愿意买单', level=3)
add_para('直接算经济账——一套系统一年15万，但帮助一家中型三甲医院每年避免的编码损失在100万以上。ROI超过6:1。再加上医保飞检罚款的风险规避（一次飞检罚款就是几十到几百万），医院的付费意愿非常明确。')

# ── 7. 竞争优势 ──
doc.add_heading('七、竞争优势', level=2)

doc.add_heading('7.1 竞品分析', level=3)

add_table_simple(['竞品', '类型', '核心弱点'], [
    ['东软望海', '传统HIS厂商', '规则引擎老旧，无AI语义理解，不能处理非结构化病历文本'],
    ['国新健康', 'DRG咨询服务商', '偏人工服务，产品化程度低，覆盖医院有限'],
    ['零氪科技', '肿瘤大数据', '只做肿瘤专科，不做全科DRG编码'],
    ['森亿智能', 'AI病历结构化', '侧重科研数据，不碰DRG付费场景'],
    ['各类小SaaS', '单一功能工具', '只做编码或只做质控，数据不互通，三套系统间信息割裂'],
])

doc.add_heading('7.2 五条核心壁垒', level=3)

add_para('1. 三合一产品形态', bold=True)
add_para('编码 + DRG分组 + 质控在一个系统里打通，数据无缝流转。传统方案是买三套系统、对接三个供应商、数据互不相通。这是市面上没有的一体化产品。')

add_para('2. AI语义理解，不是关键词匹配', bold=True)
add_para('基于LLM的编码推荐和质控检查能理解医学语义（"胸口像石头压着一样疼" → 心绞痛），而不是简单的关键词库匹配。NLP+LLM的组合能处理医生口语化、非标准化的表达。')

add_para('3. 双后端离线可用', bold=True)
add_para('Ollama本地部署 + RuleBased自动兜底。即使医院没有GPU或网络受限，核心功能依然可用。这是纯云端方案做不到的。')

add_para('4. 数据飞轮效应', bold=True)
add_para('每多一家医院使用，ICD编码数据库就更完善，质控规则就更精准。这是传统HIS厂商的静态规则引擎无法追赶的。')

add_para('5. 先发优势', bold=True)
add_para('AI + DRG + 质控一体化这个细分赛道，目前还没有成熟的商业化产品。率先进入的窗口期大约12-18个月。')

# ── 8. 比赛策略 ──
doc.add_heading('八、比赛策略', level=2)

doc.add_heading('8.1 赛道选择', level=3)
add_bullet('主攻：中国国际"互联网+"大学生创新创业大赛（教育部主办，影响力最大）')
add_bullet('第二战场："挑战杯"全国大学生创业计划竞赛（团中央主办，同一套材料可复用）')
add_bullet('备选：全国大学生生命科学竞赛（医疗AI方向垂直加分）')

doc.add_heading('8.2 评分解构（互联网+ 满分100分）', level=3)

add_table_simple(['维度', '分值', '我们要证明的'], [
    ['创新性', '30分', 'AI Agent + 自研医学知识库，不是简单调API。LLM + 向量检索 + 规则引擎三层架构'],
    ['商业性', '30分', 'DRG付费改革是国家级刚需，医院编码科是真实付费方，ROI清晰可量化（>6:1）'],
    ['团队', '20分', 'Claude全栈技术研发，有完整可运行的系统（不是PPT创业），代码可见可演示'],
    ['社会价值', '20分', '医保控费每年涉及万亿级基金使用效率，编码质量直接影响基金安全和使用公平性'],
])

doc.add_heading('8.3 路演Demo核心脚本（5分钟）', level=3)

add_table_simple(['时间段', '环节', '内容要点'], [
    ['0:00-0:30', '问题冲击', '三甲医院编码员深夜加班 + 全国10万缺口 + 编码错误率15% + 飞检罚款数据'],
    ['0:30-1:30', '产品演示', '打开码医系统 -> 导入病历 -> AI秒出ICD编码 -> 对比人工（漏2个诊断=损失8320元=一年几十万）'],
    ['1:30-3:00', '质控演示', '同一份病历的AI质控报告 -> 逐条展示缺陷 -> 传统抽查覆盖率<10% vs AI全覆盖'],
    ['3:00-4:00', '数据看板', '全院DRG运营数据：CMI评估、收入对比、高频缺陷排行 -> 管理者真正需要这些'],
    ['4:00-5:00', '商业模式+愿景', '已开发完成可演示、编码准确率目标95%+、让每一份病历都准确，让每一分医保基金都花在刀刃上'],
])

doc.add_heading('8.4 Demo的三个制胜点', level=3)
add_bullet('真人实跑：现场导入病历，AI当场出结果，不要录屏。现场感是杀伤力最大的武器。')
add_bullet('每个环节都换算成钱："漏编1个诊断编码 = DRG分组重量级变轻 = 医院损失8320元 = 一年就是几十万"——评委听钱最敏感。')
add_bullet('社会价值可视化：最后放一个计数器——"全国推广后，每年帮医保基金精准支付，节省数百亿浪费"。')

# ── 9. 团队 ──
doc.add_heading('九、团队分工', level=2)

add_table_simple(['角色', '成员', '职责'], [
    ['项目负责人', '郑诗东和（老大）', '方向决策、资源协调、商业计划书、路演PPT、路演演讲、医院合作洽谈'],
    ['技术负责人', 'Claude（AI Agent）', '全部代码开发、技术架构设计、技术白皮书编写、Demo视频制作'],
    ['医学顾问（待招募）', 'TBD', 'ICD/DRG编码规则审核、质控规则验证、医院临床试点对接'],
])

# ── 10. 下一步计划 ──
doc.add_heading('十、下一步计划', level=2)

add_table_simple(['时间', '里程碑'], [
    ['已完成', '核心功能开发：ICD编码 + DRG分组 + 质控引擎 + Pipeline + Dashboard + Admin管理 + 数据导出'],
    ['已完成', '产品打磨：ICD数据源统一(300+诊断/113手术)、Dashboard自动种子、Pipeline持久化、文件上传(.docx/.pdf)'],
    ['已完成', '竞赛基础：Git仓库初始化、README、项目介绍文档、Swagger API文档链接'],
    ['本周', 'DRG完整分组匹配 -> 商业计划书定稿 -> 路演PPT制作'],
    ['2周内', '质控规则扩充到50+条 -> 编码准确率系统测试 -> 演示脚本打磨'],
    ['1个月内', '联系1-2家医院试点 -> 收集真实试用数据 -> 迭代优化'],
    ['2-3个月', '校赛 -> 省赛 -> 根据评委反馈迭代 -> 备战国赛'],
])

# ── Footer ──
add_divider()
add_para('')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('码医 MediCode — 让每一份病历都准确，让每一分医保基金都花在刀刃上')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run('项目路径：C:\\Users\\Donghe\\Desktop\\码医-MediCode\\')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer3.add_run('更新时间：2026-05-24')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Save
output_path = r'C:\Users\Donghe\Desktop\码医-MediCode-项目介绍.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
