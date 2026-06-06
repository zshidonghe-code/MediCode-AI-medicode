"""
Generate professional dual-purpose resume .docx
Usage: python scripts/build_resume.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def set_font(run, name="微软雅黑", size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_divider(color="D0D0D0", space_before=4, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, "微软雅黑", 11.5, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
    add_divider("333333", space_before=2, space_after=6)


def bullet(text, indent=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run("• " + text)
    set_font(run, "微软雅黑", 9, color=RGBColor(0x3A, 0x3A, 0x3A))


def project_header(name, role, period):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    set_font(run, "微软雅黑", 10.5, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
    run2 = p.add_run("    " + role + "    " + period)
    set_font(run2, "微软雅黑", 8.5, color=RGBColor(0x88, 0x88, 0x88))


def metric_bullet(text, bold_words):
    """Bullet point where specified words appear in bold."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.35

    # Add bullet
    run = p.add_run("• ")
    set_font(run, "微软雅黑", 9, color=RGBColor(0x3A, 0x3A, 0x3A))

    remaining = text
    for bw in bold_words:
        if bw in remaining:
            before, after = remaining.split(bw, 1)
            if before:
                r = p.add_run(before)
                set_font(r, "微软雅黑", 9, color=RGBColor(0x3A, 0x3A, 0x3A))
            r = p.add_run(bw)
            set_font(r, "微软雅黑", 9, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
            remaining = after
    if remaining:
        r = p.add_run(remaining)
        set_font(r, "微软雅黑", 9, color=RGBColor(0x3A, 0x3A, 0x3A))


# ═══════════ HEADER ═══════════

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(2)
run = name_p.add_run("郑 诗 东 和")
set_font(run, "微软雅黑", 22, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(2)
run = contact_p.add_run("25024069@suibe.edu.cn  |  上海松江")
set_font(run, "微软雅黑", 8.5, color=RGBColor(0x77, 0x77, 0x77))

links_p = doc.add_paragraph()
links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
links_p.paragraph_format.space_after = Pt(0)
run = links_p.add_run("huanban.cloud  |  码医 MediCode")
set_font(run, "微软雅黑", 8.5, color=RGBColor(0x55, 0x55, 0x55))

add_divider("C0C0C0", 8, 4)

# ═══════════ SUMMARY ═══════════

summary_p = doc.add_paragraph()
summary_p.paragraph_format.space_before = Pt(6)
summary_p.paragraph_format.space_after = Pt(0)
summary_p.paragraph_format.line_spacing = 1.4
summary_text = (
    "大一学生创业者，名下持有两家实业公司，善于运用 AI 工具从零到一交付商业级产品。"
    "目前主导 2 个大学生创业参赛项目（家庭 SaaS / AI 医疗），覆盖产品设计、技术架构、商业运营全流程。"
    "对经济学与投资分析有浓厚兴趣，商业实践激发了对系统理论学习的强烈需求。"
    "致力于在 AI 与商业的交叉领域创造价值。"
)
run = summary_p.add_run(summary_text)
set_font(run, "微软雅黑", 9, color=RGBColor(0x44, 0x44, 0x44))

# ═══════════ PROJECTS ═══════════

section_title("项目经历")

# --- 缓伴 ---
project_header("缓伴（huanban.cloud）", "联合创始人 & 项目负责人", "2025.09 - 至今")

b1 = [
    ("借助 AI 主导产品从 0 到 1 全流程：市场调研 → 需求定义 → 技术选型 → UI 设计 → 前后端开发 → 部署上线",
     ["借助 AI", "从 0 到 1"]),
    ("项目申报主体：未名川流（上海）实业有限公司（本人担任法定代表人），中国国际大学生创新大赛参赛项目",
     ["未名川流（上海）实业有限公司", "法定代表人"]),
    ("完成全栈开发：React 前端 + Python 后端 + MySQL 数据库，适配移动端和桌面端",
     ["完成全栈开发"]),
    ("设计双角色系统（家长端/子女端），集成每日照片打卡、积分徽章激励机制、实时通讯三大核心模块",
     ["双角色系统", "三大核心模块"]),
    ("制定产品定价策略与用户增长模型，锁定高校家庭社群作为首批种子用户",
     ["定价策略", "用户增长模型"]),
    ("ICP 备案已完成终审，域名 huanban.cloud 预计两周内正式上线运营",
     ["ICP 备案已完成终审", "huanban.cloud"]),
]
for text, metrics in b1:
    metric_bullet(text, metrics)

# --- 码医 ---
project_header("码医（MediCode）", "创始人 & 项目负责人", "2025.09 - 至今")

b2 = [
    ("AI 驱动的医疗 DRG 编码与病历质控 SaaS 系统，面向医院病案科，全国大学生创业大赛参赛项目",
     ["AI 驱动", "SaaS 系统"]),
    ("自研核心引擎：ICD 编码覆盖 920 条诊断 + 611 条手术编码、CHS-DRG 1.2 分组引擎、17 条质控规则 + LLM 智能审核",
     ["920 条诊断", "611 条手术", "17 条质控规则"]),
    ("203 份真实病历基准测试：诊断 Top-1 准确率达 94.1%，编码引擎达到实用水平",
     ["94.1%"]),
    ("技术栈：FastAPI + SQLAlchemy 2.0 异步 + React 18 + TypeScript + Ant Design 5 + Docker 容器化部署",
     ["FastAPI + SQLAlchemy 2.0 异步", "React 18 + TypeScript"]),
    ("交付 26 个 RESTful API 端点、67 个自动化测试（全部通过）、商业计划书、20 页路演 PPT、5 分钟演示脚本",
     ["26 个", "67 个"]),
]
for text, metrics in b2:
    metric_bullet(text, metrics)

# --- 量化工具 ---
project_header("A 股量化分析工具", "AI 辅助开发", "2025.10 - 至今")

b3 = [
    ("实时行情监控与波段交易辅助系统，接入 A 股实时数据源",
     ["实时行情监控"]),
    ("实现 4 种形态识别算法：龙头首阴、放量突破、断板反包、N 型反转，支持参数自定义",
     ["4 种形态识别算法"]),
    ("集成技术指标计算、可视化图表输出、交易日志记录三大模块，用于投资决策学习与方法验证",
     ["三大模块"]),
]
for text, metrics in b3:
    metric_bullet(text, metrics)

# ═══════════ ENTERPRISE ═══════════

section_title("企业背景")

ent1_p = doc.add_paragraph()
ent1_p.paragraph_format.space_before = Pt(4)
ent1_p.paragraph_format.left_indent = Cm(0.3)
run = ent1_p.add_run("未名川流（上海）实业有限公司")
set_font(run, "微软雅黑", 10, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
run2 = ent1_p.add_run("    法定代表人    ")
set_font(run2, "微软雅黑", 8.5, color=RGBColor(0x88, 0x88, 0x88))
bullet("主营业务：技术服务与推广，目前作为缓伴（huanban.cloud）项目的法律实体与申报主体")

ent2_p = doc.add_paragraph()
ent2_p.paragraph_format.space_before = Pt(6)
ent2_p.paragraph_format.left_indent = Cm(0.3)
run = ent2_p.add_run("特别好看实业（上海）有限责任公司")
set_font(run, "微软雅黑", 10, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
run2 = ent2_p.add_run("    法人 / 家族运营    ")
set_font(run2, "微软雅黑", 8.5, color=RGBColor(0x88, 0x88, 0x88))
bullet("主营业务：食品批发，由家族成员参与运营，目前有在营项目，涉及供应链管理与渠道分销")

# ═══════════ SKILLS ═══════════

section_title("技术能力")

skills = [
    ("AI 应用", "大语言模型应用开发（Ollama / OpenAI）、Prompt Engineering、AI 辅助全栈开发、NLP 文本处理"),
    ("后端开发", "Python、FastAPI、Flask、SQLAlchemy 2.0 异步、RESTful API 设计"),
    ("前端开发", "React 18、TypeScript、Ant Design 5、ECharts、Vite、HTML/CSS"),
    ("数据库", "SQLite、MySQL、数据库建模与范式设计、SQLAlchemy ORM、Alembic 迁移"),
    ("DevOps", "Git / GitHub、Docker 容器化、Nginx 反向代理、CI/CD（GitHub Actions）"),
    ("商业素养", "产品定位与定价、用户增长模型设计、市场分析、商业计划书撰写、路演演讲"),
]

for cat, items in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(cat + "：")
    set_font(r1, "微软雅黑", 9, bold=True, color=RGBColor(0x2A, 0x2A, 0x2A))
    r2 = p.add_run(items)
    set_font(r2, "微软雅黑", 9, color=RGBColor(0x4A, 0x4A, 0x4A))

# ═══════════ EDUCATION ═══════════

section_title("教育背景")

edu_p = doc.add_paragraph()
edu_p.paragraph_format.space_before = Pt(2)
edu_p.paragraph_format.left_indent = Cm(0.3)
run = edu_p.add_run("上海对外经贸大学    物流管理（中澳合作）    本科 2025 级")
set_font(run, "微软雅黑", 9.5, color=RGBColor(0x2A, 0x2A, 0x2A))

course_p = doc.add_paragraph()
course_p.paragraph_format.space_before = Pt(2)
course_p.paragraph_format.left_indent = Cm(0.3)
run = course_p.add_run("主修课程：微积分 B、概率论与数理统计、微观经济学、Python 程序设计、通用英语（读写/视听说）")
set_font(run, "微软雅黑", 8.5, color=RGBColor(0x77, 0x77, 0x77))

extra_p = doc.add_paragraph()
extra_p.paragraph_format.space_before = Pt(1)
extra_p.paragraph_format.left_indent = Cm(0.3)
run = extra_p.add_run("自主修读：《中级微观经济学》（旁听，主动加深经济学理论训练）")
set_font(run, "微软雅黑", 8.5, color=RGBColor(0x77, 0x77, 0x77))

# ═══════════ COMPETITIONS ═══════════

section_title("竞赛与实践")

for item in [
    "全国大学生创业大赛（参赛中）—— 项目「码医」，完成技术开发 + 商业计划 + 路演材料",
    "中国国际大学生创新大赛（参赛中）—— 项目「缓伴」，负责产品设计与全栈开发",
]:
    bullet(item)

# ═══════════ ADDITIONAL ═══════════

section_title("其他")

for item in [
    "语言能力：具备英文技术文档阅读与编写能力",
    "个人特质：善于运用 AI 工具将想法快速落地为产品，具备独立攻克复杂问题的能力",
    "兴趣方向：创业经济学、量化投资、SaaS 商业模式、AI + 垂直行业应用",
]:
    bullet(item)

# ─── SAVE ───
output_path = "output/郑诗东和-个人简历.docx"
doc.save(output_path)
print(f"Done: {output_path}")
