# -*- coding: utf-8 -*-
"""Generate pitch pitfall checklist — honest competitive positioning for MediCode."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

LQ = '“'  # "
RQ = '”'  # "

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.4
style.paragraph_format.space_after = Pt(4)

rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), 'SimSun')
rPr.insert(0, rFonts)

# Colors
RED = RGBColor(0xc0, 0x39, 0x2b)
GREEN = RGBColor(0x1e, 0x84, 0x45)
BLUE = RGBColor(0x1a, 0x3c, 0x6e)
WINE = RGBColor(0x72, 0x2f, 0x37)
GRAY = RGBColor(0x7f, 0x8c, 0x8d)


def add_title(text, size=20, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    if color:
        run.font.color.rgb = color
    return p


def add_section(text, size=13, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    if color:
        run.font.color.rgb = color
    return p


def add_sub(text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    return p


def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_green(text):
    p = doc.add_paragraph()
    run = p.add_run('✅ ' + text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GREEN
    return p


def add_red(text):
    p = doc.add_paragraph()
    run = p.add_run('❌ ' + text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RED
    return p


def add_qa(q, a):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('Q: ' + q)
    run.font.size = Pt(10.5)
    run.font.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run('A: ' + a)
    run2.font.size = Pt(10.5)
    return p2


def add_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    return p


# ======================= CONTENT =======================

add_title('码医 MediCode 路演避坑清单', color=WINE)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('竞赛答辩 · 话术红线 · 常见质疑与标准回答')
run.font.size = Pt(9)
run.font.color.rgb = GRAY

# ============ PART 1: WHAT YOU CAN SAY ============
add_section('一、这些话可以说（真实、可验证、有数据支撑）', color=GREEN)

add_green('我们是学生团队，用开源大模型+规则引擎做了一个一体化Demo，编码+分组+质控三合一')
add_body('    → “学生团队”是护身符，不是劣势。你越坦荡，评委越宽容。')

add_green('我们用本地模型（Ollama+Qwen）而非调用云端API，数据不出医院内网——这在隐私合规场景下是有真实优势的')
add_body('    → 这是你唯一能拿出手的技术差异点，别乱用。对手用云端API处理病例=数据出医院，这是很多医院不愿意的。')

add_green('我们的Demo覆盖了从病历输入到DRG分组到质控报告的完整流程，做到了“一键式”的用户体验')
add_body('    → 说“一体化Demo”而非“一体化产品”。“Demo”这个词是你最好的保护。')

add_green('我们在小样本（203份）上的编码Top-1准确率为94.1%，这是实验室环境下的基线数据')
add_body('    → 关键词：“小样本”“实验室环境”“基线”。不要拿这个数字去对标讯飞的91%(百万级训练集)。')

# ============ PART 2: WHAT YOU MUST NEVER SAY ============
add_section('二、这些话绝对不能说（一说就翻车）', color=RED)

add_red('“市场上没有同类产品” / “我们是蓝海独家”')
add_body('    → 评委里有医院信息科主任，他们每天在用艾登/讯飞/侠医。你一说这句话，后面所有内容都不信了。')

add_red('“大公司不屑做、小公司做不了”')
add_body('    → IBM的故事可以用来回答“为什么你们敢做”，但不能拿来说“没人做”。事实是一堆人在做。把这句话从所有材料里删掉。')

add_red('“编码准确率94.1%”（没有前置限定词）')
add_body('    → 必须加“在小样本实验室环境下”。否则评委一句“你这个数据是多大样本测的？”，直接翻车。')

add_red('“我们的系统可以替代编码员”')
add_body('    → 编码员是国家职业资格，有法律责任的。你只能说“辅助”“提效”“第一道筛查”。')

add_red('“我们的系统可以直接落地医院”')
add_body('    → 你连等保三级、HIS对接、高并发、安全审计都没做。说“原型验证”而非“商用落地”。')

add_red('准确率对标友商“我们94%比XXX高”')
add_body('    → 别人是百万级训练集+真实医院环境，你是203份干净样本。这不是对比，是自杀。')

# ============ PART 3: TOUGH QUESTIONS ============
add_section('三、评委常问的刁钻问题 + 标准回答', color=BLUE)

add_rule()

add_qa(
    '你们训练数据只有203份病历，讯飞、艾登是百万级。你凭什么说你的方案可行？',
    '“数据量确实是我们目前最大的短板。203份的测试结果只能证明技术路线可行，不能证明产品成熟。但我们的思路和已有产品不同：他们用百万级数据训练一个通用模型，表面上很强，但换一家医院就要重新调参。我们走的是“规则引擎+小模型”路线，规则是硬编码的，可解释、可调整，不依赖海量数据。这个思路在数据稀缺场景下反而是优势。”'
)

add_rule()

add_qa(
    '你们团队没有临床医生，凭什么做医疗AI？',
    '“这是我们目前最大的痛点，我不回避。但两点：第一，ICD编码本质上是分类任务，不是诊断任务——我们不需要学会看病，只需要学会把病历里的关键词映射到920个编码上。第二，我们正在积极寻找医学顾问（已与XX医院取得联系），这是我们接下来的第一优先级。”'
)

add_rule()

add_qa(
    '市场上已经有艾登、讯飞、侠医、百度灵医，你们凭什么竞争？',
    '“他们做得很好，我们不试图否认这一点。但我们的切入点不同：第一，我们走本地模型路线，数据不出医院内网——对于对数据安全敏感的医院，这是刚需。第二，我们的规则引擎是白盒化的，医院可以自己调整编码规则——而艾登、讯飞的模型是黑盒。第三，我们的目标市场是他们覆盖不到的长尾医院——二级医院买不起艾登的年费，但又需要基本的编码辅助。”'
)

add_rule()

add_qa(
    '你的系统在真实医院环境测试过吗？准确率多少？',
    '“没有。我们目前只在实验室环境下做了基线测试，真实医院环境下的准确率会低很多。这正是我们接下来要做的事——找一家愿意配合试点的医院，拿真实数据来赟。”'
)

add_rule()

add_qa(
    '你们这个项目能赚钱吗？商业模式是什么？',
    '“目前没有任何付费用户，这也是事实。但这个问题的另一面是：一个和艾登、讯飞完全同质化的商业模式确实赚不到钱。但我们的价值不是“另一个艾登+”，而是“可私有化部署的轻量级编码工具”，这个定位在二级医院和民营医院市场是有真实需求的。”'
)

add_rule()

add_qa(
    '你是一个人做的吗？团队到底有多少人？',
    '“技术开发确实是我一个人完成的，但我们正在扩展团队——目前正在联系医学顾问、寻找商科同学负责商业计划。一个人写完所有代码是事实，但正是因为用了AI开发工具，效率才能这么高。这本身就证明了我们项目核心理念的可行性——AI可以让一个人做以前十个人的事。”'
)

# ============ PART 4: POSITIONING ============
add_section('四、路演里如何定位自己（不死的说法）', color=BLUE)

add_sub('“我们不是‘替代者’，我们是‘补充者’”')
add_body(
    '艾登/讯飞做三甲医院，年费几十万，买不起的二级/民营医院是我们的空间。'
    '这个定位是真实的，不是编的。'
)

add_sub('“我们不是‘下一个艾登’，我们是‘可私有化部署的轻量级编码工具’”')
add_body(
    '别人用云端API、黑盒模型、不可解释；'
    '我们用本地模型、白盒规则、医院可自行调整。'
    '这是真实的差异，不是硬凹的卖点。'
)

add_sub('“我们是学生项目，不是商业产品”')
add_body(
    '所有不足——数据少、没有临床医生、没有付费用户——只要你主动说“我们是学生项目”，'
    '这些都变成了“当然的不足”而非“致命的硬伤”。'
    '“学生项目”四个字是你最强的护盾，别不用。'
)

# ============ PART 5: COMPETITOR LANDSCAPE (one glance) ============
add_section('五、竞品一览（路演时别被问倒）', color=BLUE)

comp_table = doc.add_table(rows=7, cols=5, style='Light Grid Accent 1')
comp_table.autofit = True

headers = ['产品', '核心能力', '规模', '价格层级', '对我们的威胁']
for i, h in enumerate(headers):
    comp_table.rows[0].cells[i].text = h
    for pp in comp_table.rows[0].cells[i].paragraphs:
        for r in pp.runs:
            r.font.bold = True
            r.font.size = Pt(8)

comp_data = [
    ['艾登科技', 'AI编码+DRG+质控，百万级训练', '近4000家医院', '几十万/年', '✅行业头部，不要正面对比'],
    ['讯飞医疗', 'NLP强，编码+质控+分组', '三甲主', '几十万/年', '口碑好，泛化能力强'],
    ['侠医软件', '编码+质控+分组', '47家医院', '中等', '小体量、灵活，有真实回款案例'],
    ['百度灵医', 'CDSS+DRG+编码', '600+医院', '大厂定价', '资源雄厚，但非核心业务'],
    ['传统HIS厂商', '编码模块（关键词匹配）', '千家级', '捆绑免费', '市场占有率极高，但AI弱'],
    ['码医 MediCode', '本地小模型+白盒规则+一体化Demo', '学生项目', '免费/低价', '—'],
]
for i, row_data in enumerate(comp_data):
    for j, val in enumerate(row_data):
        comp_table.rows[i+1].cells[j].text = val
        for pp in comp_table.rows[i+1].cells[j].paragraphs:
            for r in pp.runs:
                r.font.size = Pt(7.5)

# Last row highlight
for cell in comp_table.rows[6].cells:
    for pp in cell.paragraphs:
        for r in pp.runs:
            r.font.bold = True

add_body('')  # spacer

# ============ PART 6: IF YOU CAN ONLY REMEMBER 3 THINGS ============
add_section('六、只能记住三件事', color=WINE)

add_bold = lambda t, r: None  # placeholder

items = [
    ('1. 永远先说“我们是学生项目”。',
     '这四个字把所有硬伤变成了“现阶段当然的不足”。'),
    ('2. 永远不要说“我们是唯一的”。',
     '承认市场上有很多人在做，然后说你的差异点。评委尊重诚实，鄙视吹牛。'),
    ('3. 永远不要拿小样本数据对标商用产品。',
     '你的94.1%和讯飞的91%不是同一个量纲的数字。说清楚“小样本实验室”这个前置条件。'),
]
for bold, rest in items:
    p = doc.add_paragraph()
    r1 = p.add_run(bold)
    r1.font.size = Pt(10.5)
    r1.font.bold = True
    r2 = p.add_run(rest)
    r2.font.size = Pt(10.5)

# Footer
footer = doc.add_paragraph()
footer.paragraph_format.space_before = Pt(16)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    '这份清单的目的不是让你少说错话，是让你知道自己真实站在哪。'
    '知道自己的边界在哪的人，比不知道自己有多弱的人，更容易赢得信任。'
)
run.font.size = Pt(8)
run.font.color.rgb = GRAY
run.font.italic = True

# Save
output_path = r'C:\Users\Donghe\Desktop\码医-路演避坑清单.docx'
doc.save(output_path)
print(f'Done: {output_path}')
