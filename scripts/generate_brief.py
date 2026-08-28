# -*- coding: utf-8 -*-
"""Generate concise project brief for middleman to evaluate and introduce."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

LQ = '“'  # "
RQ = '”'  # "

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), 'SimSun')
rPr.insert(0, rFonts)


def add_heading(text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    return p


def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_bullet(text):
    p = doc.add_paragraph()
    run = p.add_run('• ' + text)
    run.font.size = Pt(11)
    return p


def add_bold_body(bold_part, rest):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_part)
    r1.font.size = Pt(11)
    r1.font.bold = True
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)
    return p


# ===================== CONTENT =====================

# Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run('码医 MediCode — 项目说明')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run('供引荐人评估使用 · ' + datetime.date.today().strftime('%Y-%m-%d'))
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x8f, 0x8a, 0x80)

# 1. What
add_heading('一、项目是什么')
add_body(
    '码医 MediCode 是一个 AI 医疗系统，核心功能就三个：'
)
add_bullet('用 AI 自动给住院病历打 ICD-10 诊断编码')
add_bullet('自动做 DRG 付费分组（医保直接按这个付钱）')
add_bullet('自动检查病历质量（完整性、逻辑性、编码规范）')
add_body(
    '一句话：帮医院不被医保扣钱，帮医保基金不被浪费。目前市场上没有同时做"编码+分组+质控"一体化产品。'
)

# 2. Current status
add_heading('二、项目现状')
add_bullet('后端 26 个 API 端点，前端 9 个页面，完整可演示')
add_bullet('920 条 ICD-10 诊断编码 + 611 条手术编码 + CHS-DRG 1.2 分组方案')
add_bullet('在 4 例合成病历上完成规则模式验证：主诊断匹配 2/4，手术编码匹配 2/4')
add_bullet('技术栈：Python FastAPI + React + Ollama 本地大模型')
add_bullet('参赛：中国国际大学生创新大赛（原"互联网+"），目标全国第一')

# 3. Why not big companies?
add_heading('三、为什么大公司不做这个')
add_body(
    '这是一个很自然会问的问题——如果这件事有价值，百度、腾讯、东软、卫宁这些巨头为什么没做？'
    '答案不是"他们没看到"，而是"他们要么做不了，要么不划算"。以下是真实的行业格局：'
)

# Case 1: IBM Watson Health
add_bold_body('案例一：IBM Watson Health —— 50亿美元的教训', '')
add_body(
    'IBM在2015年成立Watson Health，先后砸了50多亿美元收购Truven、Merge Healthcare等多家医疗数据公司，'
    '最多时有7000名员工。他们声称要用AI改变癌症治疗和临床决策。结果呢？'
)
add_bullet('与MD Anderson癌症中心的合作在2017年终止，投入6200万美元后没有产出任何可用产品，一份措辞严厉的审计报告随之曝光')
add_bullet('多项独立研究显示，Watson在临床任务中的准确率不足50%，经常推荐不相关或不可用的治疗方案')
add_bullet('2017年一份IBM内部演示文档中，一位合作医生直接称其产品为' + LQ + '垃圾' + RQ)
add_bullet('2022年1月，IBM将Watson Health核心资产以约10亿美元贱卖——投资50亿，回收10亿，净亏40亿')
add_body(
    '教训：即便是地球上最顶尖的AI公司，砸了50亿美元和7年时间，也没有在"把医学知识变成AI规则"这件事上成功。因为他们试图用通用AI解决一个需要深度医学知识工程的问题——这正是码医从第一天就绕开的坑：我们不做"万能诊断"，只做ICD编码这一个窄而深的切口。'
)

# Case 2: Google Health
add_bold_body('案例二：Google Health —— 两年散伙', '')
add_body(
    '谷歌在2018年高调组建Google Health部门，从全美顶级医疗机构Geisinger挖来David Feinberg担任副总裁。'
    '三年后的2021年，Google Health作为一个独立部门被解散，Feinberg离职。核心原因：'
    '谷歌擅长的是"搜索全世界的网页"，不是"理解一份病历里的上下文"。'
    '医疗数据的碎片化、隐私合规（HIPAA）、医院采购周期——都跟谷歌的"快速迭代"文化水土不服。'
)

# Case 3: BAT
add_bold_body('案例三：百度、阿里、腾讯的医疗AI —— 全在另一个赛道', '')
add_body(
    '国内三大巨头的医疗AI布局看似热闹，但没有一家在做"AI辅助ICD编码"这件事：'
)
add_bullet('腾讯觅影：聚焦影像筛查（肿瘤CT/MRI）+ 医保控费报表，不是编码辅助')
add_bullet('百度灵医：做临床决策支持（CDSS）+ DRG控费看板，编码能力停留在规则匹配层面')
add_bullet('阿里健康AI：核心是医疗影像分析（CT/MRI病灶标注）+ 医药电商，完全不在编码赛道')
add_body(
    '为什么？因为ICD编码太"脏"了——它不是一个光鲜的神经网络问题，'
    '而是一个需要理解920条诊断编码规则、每家医院的编码习惯、临床上下文判断的知识工程。'
    '大厂做这个，相当于用大炮打蚊子，投入产出比算不过账。'
)

# Case 4: 传统HIS厂商
add_bold_body('案例四：东软、卫宁等HIS厂商 —— 有编码模块，无AI能力', '')
add_body(
    '国内医疗IT三巨头东软、卫宁、创业慧康都有DRG模块，但它们的产品有三个硬伤：'
)
add_bullet('编码方式：基于关键词模板匹配，不是NLP+LLM语义理解——' + LQ + '肺炎' + RQ + '能匹配到，' + LQ + '右下肺斑片状浸润影，考虑社区获得性肺炎' + RQ + '就匹配不到')
add_bullet('产品定位：DRG模块是HIS系统的一个附属功能，目的是帮医院"生成医保局要求的报表"，不是"帮编码员提高准确率和效率"')
add_bullet('迭代速度：HIS厂商的更新周期以年计，而DRG政策和编码规范每年都在变——医院等不起')
add_body(
    '所以现状是：大厂有AI但嫌市场小，HIS厂商有渠道但缺AI，两者之间留下了一个清晰的空白地带。'
)

add_body(
    '总结：IBM砸了50亿证明了"大公司做不好医疗知识工程"；Google证明了"互联网文化≠医疗文化"；'
    'BAT选了更容易赚钱的影像和药品赛道；HIS厂商守着旧技术吃老本。'
    '这个市场不会被大公司抢走——不是因为他们没能力，而是因为他们不会为一个"只有45亿"的市场破自己的规模化逻辑。'
    '而对一个聚焦的团队来说，这恰恰是"大公司不屑做、小公司做不了"的蓝海缝隙。'
)

# 4. Why her
add_heading('四、为什么想找金花老师')
add_body(
    '不是随便找的。金花老师的三个标签，跟这个项目精确对齐：'
)
add_bold_body('全科医学临床质控中心秘书：', '码医核心功能就是病历质控，她正是质控标准的制定者。她的日常判断可以直接变成 AI 的规则逻辑。')
add_bold_body('多病共存诊疗决策研究：', 'DRG 编码最核心的难点就是合并症/并发症（CC/MCC）识别——这恰好是她的研究方向。')
add_bold_body('同济大学附属医院 · 博导：', '学术身份意味着她理解什么样的产出能发表、能申报课题，合作不止于挂名。')

# 5. What we want
add_heading('五、希望的合作方式')
add_bullet('指导老师署名（竞赛报名材料中作为医学顾问）')
add_bullet('审阅 AI 编码逻辑和质控规则的临床合理性（1-2 次）')
add_bullet('如有余力，路演答辩前给予方向性指导')
add_body(
    '核心角色：确保项目在医学上"不说外行话、不做外行事"。时间投入克制，但产出导向明确——'
    '竞赛获奖（指导老师署名）+ 后续论文发表（通讯作者）+ 课题申报（预研基础）。'
)

# 6. Team
add_heading('六、团队')
add_bullet('郑诗东和（项目负责人）：上海对外经贸大学大一，全栈开发 + 路演答辩')
add_bullet('利用 AI 开发工具完成全部代码开发，规则模式验证结果已归档')
add_bullet('已产出：商业计划书、执行摘要、Demo 录屏、20 页路演 PPT')

# 7. Quick facts
add_heading('七、关键数字')
data_points = [
    ('编码验证：', '主诊断 2/4、手术 2/4（4 例规则模式）'),
    ('编码知识库：', '920 条 ICD-10 诊断 + 611 条 ICD-9-CM-3 手术'),
    ('质控规则：', '17 条核心规则 + NLP + LLM 语义检查'),
    ('市场痛点：', '全国编码员缺口超 10 万人，人工错误率 10-15%'),
    ('竞赛目标：', '中国国际大学生创新大赛 全国第一'),
]
for bold, rest in data_points:
    add_bold_body(bold, rest)

# Footer note
footer = doc.add_paragraph()
footer.paragraph_format.space_before = Pt(16)
run = footer.add_run(
    '附：商业计划书、执行摘要、系统 Demo 录屏。如需当面演示，随时可以。\n'
    '联系人：郑诗东和 / 15800565959 / 1975790036@qq.com'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x8f, 0x8a, 0x80)

# Save
output_path = r'C:\Users\Donghe\Desktop\码医项目说明-供引荐人.docx'
doc.save(output_path)
print(f'Done: {output_path}')
